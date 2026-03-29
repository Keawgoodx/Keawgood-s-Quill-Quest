"""
╔══════════════════════════════════════════════════════════════════════════════╗
║          Keawgood Universe  v5.0  —  All-in-One Hub                        ║
║  รวมทุกโปรแกรมไว้ในไฟล์เดียว ทุกโปรแกรมใช้งานได้ 100%                  ║
║  ✦ NEW: Novel Checker (เช็คตอนนิยายครบ) — รวมจาก novel_checker.py       ║
║  ✦ UX/UI ยกระดับขึ้นสู่ระดับ Expert Design ทุก Module                    ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

# ══════════════════════════════════════════════════════════
#  SHARED IMPORTS
# ══════════════════════════════════════════════════════════
import os
import re
import shutil
import threading
import concurrent.futures
import json
import time
import random
from pathlib import Path
from datetime import datetime
from urllib.parse import urljoin

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import tkinter.font as tkfont

import customtkinter as ctk

# ══════════════════════════════════════════════════════════════════════════════
#
#  ██████╗ ██╗   ██╗    ██╗  ██╗███████╗ █████╗ ██╗    ██╗ ██████╗  ██████╗  ██████╗ ██████╗
#  ██╔══██╗╚██╗ ██╔╝    ██║ ██╔╝██╔════╝██╔══██╗██║    ██║██╔════╝ ██╔═══██╗██╔═══██╗██╔══██╗
#  ██████╔╝ ╚████╔╝     █████╔╝ █████╗  ███████║██║ █╗ ██║██║  ███╗██║   ██║██║   ██║██║  ██║
#  ██╔══██╗  ╚██╔╝      ██╔═██╗ ██╔══╝  ██╔══██║██║███╗██║██║   ██║██║   ██║██║   ██║██║  ██║
#  ██████╔╝   ██║       ██║  ██╗███████╗██║  ██║╚███╔███╔╝╚██████╔╝╚██████╔╝╚██████╔╝██████╔╝
#  ╚═════╝    ╚═╝       ╚═╝  ╚═╝╚══════╝╚═╝  ╚═╝ ╚══╝╚══╝  ╚═════╝  ╚═════╝  ╚═════╝╚═════╝
#
#  by_keawgood.py — Thai Novel File Manager
# ══════════════════════════════════════════════════════════════════════════════

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import fitz  # PyMuPDF
    PDF_OK = True
except ImportError:
    PDF_OK = False

# ── Colors ──
BK_ACCENT   = ("#6366F1", "#7C6BFF")
BK_ACCENT2  = ("#4F46E5", "#A78BFA")
BK_BG_DARK  = ("#F1F5F9", "#0F0F1A")
BK_BG_CARD  = ("#FFFFFF", "#1A1A2E")
BK_BG_INPUT = ("#E2E8F0", "#252540")
BK_FG_TEXT  = ("#1E293B", "#E2E8F0")
BK_FG_MUTED = ("#64748B", "#94A3B8")
BK_SUCCESS  = ("#10B981", "#34D399")
BK_ERROR    = ("#EF4444", "#F87171")
BK_WARN     = ("#F59E0B", "#FBBF24")

CHAPTER_PATTERN = re.compile(r'^(?:ตอนที่|ตอน|บทที่|บท|Chapter)\s*\d+', re.IGNORECASE)


def bk_natural_key(name: str):
    nums = re.findall(r'\d+', name)
    return int(nums[0]) if nums else 0


def bk_read_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        if not PDF_OK:
            raise ImportError("PyMuPDF ไม่ได้ติดตั้ง — pip install PyMuPDF")
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    if ext == ".docx":
        if not DOCX_OK:
            raise ImportError("python-docx ไม่ได้ติดตั้ง — pip install python-docx")
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs)
    for enc in ("utf-8-sig", "utf-8", "tis-620", "cp874"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"อ่านไฟล์ไม่ได้: {path}")


def bk_write_text(path: str, text: str):
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8-sig") as f:
        f.write(text)


def bk_write_docx(path: str, text: str):
    if not DOCX_OK:
        raise ImportError("python-docx ไม่ได้ติดตั้ง")
    doc = DocxDocument()
    for line in text.splitlines():
        line = line.strip('\u200b\ufeff \t')
        if CHAPTER_PATTERN.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        else:
            doc.add_paragraph(line)
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc.save(path)


def bk_clean_text(text: str) -> str:
    text = text.replace('\u200b', '').replace('\ufeff', '')
    lines = [l.rstrip() for l in text.splitlines()]
    return "\n".join(lines)


def bk_safe_filename(name: str, max_len: int = 100) -> str:
    name = re.sub(r'[\\/*?:"<>|]', "", name).strip()
    return name[:max_len].strip() or "unnamed"


def bk_list_files(folder: str, exts=(".txt", ".md", ".docx", ".pdf")):
    files = [
        f for f in os.listdir(folder)
        if os.path.isfile(os.path.join(folder, f))
        and os.path.splitext(f)[1].lower() in exts
    ]
    files.sort(key=bk_natural_key)
    return files


def bk_merge_files(src_folder, out_file, log):
    files = bk_list_files(src_folder)
    if not files:
        log("❌ ไม่พบไฟล์ในโฟลเดอร์"); return
    log(f"📂 พบไฟล์ {len(files)} ไฟล์ กำลังรวม...")
    chunks = []
    for fn in files:
        path = os.path.join(src_folder, fn)
        if os.path.abspath(path) == os.path.abspath(out_file):
            continue
        try:
            chunks.append(bk_clean_text(bk_read_text(path)))
            log(f"   ✅ {fn}")
        except Exception as e:
            log(f"   ⚠️ {fn} — {e}")
    combined = "\n\n---\n\n".join(chunks)
    if os.path.splitext(out_file)[1].lower() == ".docx":
        bk_write_docx(out_file, combined)
    else:
        bk_write_text(out_file, combined)
    log(f"\n🎉 รวมไฟล์เสร็จแล้ว → {out_file}")


def bk_split_to_chapters(src_path_or_folder, out_folder, out_ext, is_folder, prefix, include_title, log):
    """
    ใช้เลขตอนจากบรรทัดแรกของแต่ละตอน (เช่น "ตอนที่ 401") แทนการนับลำดับ
    include_title=True  -> ตอนที่ 401 ตกงานแล้ว ต้องดูแลพี่สะใภ้.txt
    include_title=False -> ตอนที่ 401.txt
    """
    paths = ([os.path.join(src_path_or_folder, f) for f in bk_list_files(src_path_or_folder)]
             if is_folder else [src_path_or_folder])
    os.makedirs(out_folder, exist_ok=True)

    total_files = 0

    for path in paths:
        log(f"📖 อ่านไฟล์เพื่อแยกตอน: {os.path.basename(path)}")
        try:
            raw = bk_clean_text(bk_read_text(path))
        except Exception as e:
            log(f"   ❌ {e}"); continue

        lines = raw.splitlines()
        chapter_first_line = ""
        chapter_lines = []

        def save_ch(first_line, content_lines):
            nonlocal total_files
            if not content_lines: return
            text = "\n".join(content_lines).strip()
            if not text: return

            p = prefix.strip()

            # ดึงเลขตอนจากบรรทัดแรกของตอนนั้นๆ เช่น "ตอนที่ 401 ..." → 401
            chapter_num_str = ""
            m_num = re.match(r'^(?:ตอนที่|ตอน|บทที่|บท|Chapter)\s*(\d+)', first_line, re.IGNORECASE)
            if m_num:
                chapter_num_str = m_num.group(1)  # เลขจริงจากเนื้อหา เช่น "401"

            # สร้าง num_part โดยแทน [n] ด้วยเลขจริง
            if p:
                if "[n]" in p:
                    num_part = p.replace("[n]", chapter_num_str) if chapter_num_str else p.replace("[n]", "?")
                else:
                    num_part = f"{p} {chapter_num_str}" if chapter_num_str else p
            else:
                num_part = chapter_num_str or "?"

            # ดึงชื่อตอน (ส่วนหลังเลข) เช่น "ตอนที่ 401 ตกงานแล้ว" → "ตกงานแล้ว"
            if include_title and first_line:
                m_title = re.match(r'^(?:ตอนที่|ตอน|บทที่|บท|Chapter)\s*\d+\s*(.*)', first_line, re.IGNORECASE)
                chapter_name = m_title.group(1).strip() if m_title else ""
                # ถ้าไม่มีชื่อตอนหลังเลข ให้ใช้ first_line ทั้งหมด (กรณีชื่อตอนไม่ขึ้นต้นด้วย pattern)
                if not chapter_name and not m_num:
                    chapter_name = first_line.strip()
                fname = bk_safe_filename(f"{num_part} {chapter_name}") if chapter_name else bk_safe_filename(num_part)
            else:
                fname = bk_safe_filename(num_part)

            fname += (out_ext or ".txt")
            fpath = os.path.join(out_folder, fname)

            if out_ext == ".docx":
                bk_write_docx(fpath, text)
            else:
                bk_write_text(fpath, text)

            total_files += 1
            log(f"   💾 บันทึก: {fname}")

        for line in lines:
            cl = line.strip('\u200b\ufeff \t\r')
            if not cl:
                chapter_lines.append(""); continue
            if CHAPTER_PATTERN.match(cl):
                save_ch(chapter_first_line, chapter_lines)
                chapter_first_line = cl
                chapter_lines = [cl]
                log(f"   🔖 เจอตอน: {cl[:60]}")
            else:
                chapter_lines.append(cl)
        save_ch(chapter_first_line, chapter_lines)

    log(f"\n🎉 แยกตอนเสร็จสิ้น! ได้ทั้งหมด {total_files} ไฟล์ → {out_folder}")


def bk_batch_merge_files(src_folder, out_folder, batch_size_str, prefix, out_ext, log):
    try:
        batch_size = int(batch_size_str)
        if batch_size <= 0: raise ValueError
    except ValueError:
        log("❌ กรุณาระบุจำนวนไฟล์ต่อ 1 กลุ่มเป็นตัวเลขที่มากกว่า 0"); return

    raw_files = bk_list_files(src_folder)
    files = []
    
    for f in raw_files:
        if " - " in f:
            continue
        if not re.search(r'\d+', f):
            continue
        files.append(f)

    if not files:
        log("❌ ไม่พบไฟล์ตอนย่อยในโฟลเดอร์ต้นทาง (เช็คว่ามีตัวเลขในชื่อและไม่มีเครื่องหมาย '-')"); return
        
    os.makedirs(out_folder, exist_ok=True)
    total_groups = 0
    
    for i in range(0, len(files), batch_size):
        batch = files[i:i+batch_size]
        
        start_num = bk_natural_key(batch[0])
        end_num = bk_natural_key(batch[-1])
        
        chunks = []
        for fn in batch:
            path = os.path.join(src_folder, fn)
            log(f"📖 กำลังรวมไฟล์: {fn}")
            try:
                chunks.append(bk_clean_text(bk_read_text(path)))
            except Exception as e:
                log(f"   ❌ ข้ามไฟล์ {fn} เนื่องจาก: {e}")
                
        if not chunks: continue
        
        combined = "\n\n---\n\n".join(chunks)
        
        p = prefix.strip()
        if p:
            fname = f"{p} {start_num:03d} - {end_num:03d}{out_ext}"
        else:
            fname = f"{start_num:03d}-{end_num:03d}{out_ext}"
            
        fpath = os.path.join(out_folder, fname)
        
        if out_ext == ".docx":
            bk_write_docx(fpath, combined)
        else:
            bk_write_text(fpath, combined)
            
        total_groups += 1
        log(f"   📦 บันทึกกลุ่มไฟล์สำเร็จ: {fname}")
        
    log(f"\n🎉 รวมกลุ่มเสร็จสิ้น! ได้ผลลัพธ์ทั้งหมด {total_groups} ไฟล์ → {out_folder}")


BK_CONVERSIONS = [
    (".txt", ".md",   "TXT → MD"),
    (".md",  ".txt",  "MD → TXT"),
    (".docx",".txt",  "DOCX → TXT"),
    (".txt", ".docx", "TXT → DOCX"),
]


def bk_convert_files(src_folder, out_folder, from_ext, to_ext, log):
    files = [f for f in os.listdir(src_folder)
             if f.lower().endswith(from_ext)
             and os.path.isfile(os.path.join(src_folder, f))]
    if not files:
        log(f"❌ ไม่พบไฟล์ {from_ext} ในโฟลเดอร์"); return
    os.makedirs(out_folder, exist_ok=True)
    count = 0
    for fn in files:
        src = os.path.join(src_folder, fn)
        base = os.path.splitext(fn)[0]
        try:
            if from_ext == ".docx" and to_ext == ".txt":
                bk_write_text(os.path.join(out_folder, base + to_ext), bk_clean_text(bk_read_text(src)))
            elif from_ext == ".txt" and to_ext == ".docx":
                bk_write_docx(os.path.join(out_folder, base + to_ext), bk_clean_text(bk_read_text(src)))
            else:
                shutil.copy2(src, os.path.join(out_folder, base + to_ext))
            count += 1
            log(f"   ✅ {fn} → {base + to_ext}")
        except Exception as e:
            log(f"   ❌ {fn}: {e}")
    log(f"\n🎉 แปลงไฟล์เสร็จ {count} ไฟล์ → {out_folder}")


# ── UI Widgets for ByKeawgood ──

class BK_LogBox(ctk.CTkTextbox):
    def __init__(self, master, **kw):
        super().__init__(master, font=("Consolas", 12), text_color=BK_FG_TEXT,
                         fg_color=BK_BG_INPUT, corner_radius=10, wrap="word", **kw)
        self.configure(state="disabled")

    def write(self, msg: str):
        self.configure(state="normal")
        self.insert("end", msg + "\n")
        self.see("end")
        self.configure(state="disabled")
        self.update_idletasks()

    def clear(self):
        self.configure(state="normal")
        self.delete("1.0", "end")
        self.configure(state="disabled")


def bk_row(parent, label, var, browse_cmd, btn_label="📂 เลือก"):
    frame = ctk.CTkFrame(parent, fg_color="transparent")
    frame.pack(fill="x", pady=4)
    ctk.CTkLabel(frame, text=label, font=("Kanit", 13), text_color=BK_FG_MUTED,
                 width=160, anchor="w").pack(side="left")
    ctk.CTkEntry(frame, textvariable=var, font=("Kanit", 12),
                 fg_color=BK_BG_INPUT, border_color=BK_ACCENT,
                 corner_radius=8).pack(side="left", fill="x", expand=True, padx=(6, 6))
    ctk.CTkButton(frame, text=btn_label, command=browse_cmd,
                  fg_color=BK_ACCENT, hover_color=BK_ACCENT2,
                  font=("Kanit", 12), width=100, corner_radius=8).pack(side="left")


def bk_section_title(parent, text):
    ctk.CTkLabel(parent, text=text, font=("Kanit", 17, "bold"),
                 text_color=BK_ACCENT2).pack(anchor="w", pady=(14, 2))
    ctk.CTkFrame(parent, height=2, fg_color=BK_ACCENT, corner_radius=1).pack(fill="x", pady=(0, 10))


def bk_run_btn(parent, text, cmd):
    return ctk.CTkButton(parent, text=text, command=cmd,
                         fg_color=BK_ACCENT, hover_color=BK_ACCENT2,
                         font=("Kanit", 14, "bold"), height=44, corner_radius=10)


class BK_MergeTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var = ctk.StringVar()
        self.out_var = ctk.StringVar()
        self.out_name_var = ctk.StringVar(value="รวมไฟล์ทั้งหมด.txt")
        self._build()

    def _build(self):
        bk_section_title(self, "  🗂️  รวมไฟล์ทั้งหมดเป็นไฟล์เดียว")
        bk_row(self, "โฟลเดอร์ต้นทาง:", self.src_var,
               lambda: self.src_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์ต้นทาง")))
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์ปลายทาง")))
        row = ctk.CTkFrame(self, fg_color="transparent")
        row.pack(fill="x", pady=4)
        ctk.CTkLabel(row, text="ชื่อไฟล์ผลลัพธ์:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        ctk.CTkEntry(row, textvariable=self.out_name_var, font=("Kanit", 12),
                     fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8
                     ).pack(side="left", fill="x", expand=True, padx=6)
        bk_run_btn(self, "▶  เริ่มรวมไฟล์", self._run).pack(pady=18)

    def _run(self):
        src = self.src_var.get().strip()
        out_dir = self.out_var.get().strip()
        name = self.out_name_var.get().strip()
        if not src or not out_dir or not name:
            messagebox.showwarning("By Keawgood", "กรุณากรอกข้อมูลให้ครบก่อนนะครับ"); return
        self.log.clear()
        threading.Thread(target=bk_merge_files,
                         args=(src, os.path.join(out_dir, name), self.log.write),
                         daemon=True).start()


class BK_SplitTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var         = ctk.StringVar()
        self.out_var         = ctk.StringVar()
        self.mode_var        = ctk.StringVar(value="file")
        self.prefix_var      = ctk.StringVar(value="ตอนที่ [n]")
        self.ext_var         = ctk.StringVar(value=".txt")
        self.include_title   = ctk.BooleanVar(value=True)   # ✅ เลือกว่าจะรวมชื่อตอนหรือไม่
        self._build()

    def _build(self):
        bk_section_title(self, "  ✂️  แยกไฟล์ (ทีละ 1 ตอน)")

        # โหมดต้นทาง
        mode_f = ctk.CTkFrame(self, fg_color="transparent")
        mode_f.pack(fill="x", pady=6)
        ctk.CTkLabel(mode_f, text="โหมด:", font=("Kanit", 13), text_color=BK_FG_MUTED,
                     width=160, anchor="w").pack(side="left")
        ctk.CTkRadioButton(mode_f, text="ไฟล์เดียว (ไฟล์รวมหลายตอน)", variable=self.mode_var, value="file",
                           font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=10)
        ctk.CTkRadioButton(mode_f, text="ทั้งโฟลเดอร์", variable=self.mode_var, value="folder",
                           font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=10)

        bk_row(self, "ต้นทาง:", self.src_var, self._browse_src)
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์สำหรับเซฟไฟล์ที่แยกแล้ว")))

        # คำนำหน้าไฟล์
        opt_f = ctk.CTkFrame(self, fg_color="transparent")
        opt_f.pack(fill="x", pady=4)
        ctk.CTkLabel(opt_f, text="คำนำหน้า ([n] = เลขตอน):", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        ctk.CTkEntry(opt_f, textvariable=self.prefix_var, font=("Kanit", 12),
                     fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8
                     ).pack(side="left", fill="x", expand=True, padx=(6, 0))

        # ✅ Checkbox: ต้องการรวมชื่อตอนหรือไม่
        name_f = ctk.CTkFrame(self, fg_color="transparent")
        name_f.pack(fill="x", pady=(2, 6))
        ctk.CTkLabel(name_f, text="รูปแบบชื่อไฟล์:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        ctk.CTkRadioButton(
            name_f,
            text="ตอนที่ [n]  ชื่อตอน   (เช่น  ตอนที่ 1 ตกงานแล้ว ต้องดูแลพี่สะใภ้.txt)",
            variable=self.include_title, value=True,
            font=("Kanit", 12), fg_color=BK_ACCENT
        ).pack(side="left", padx=(0, 16))
        ctk.CTkRadioButton(
            name_f,
            text="ตอนที่ [n]  เท่านั้น   (เช่น  ตอนที่ 1.txt)",
            variable=self.include_title, value=False,
            font=("Kanit", 12), fg_color=BK_ACCENT
        ).pack(side="left")

        # นามสกุลไฟล์
        ext_f = ctk.CTkFrame(self, fg_color="transparent")
        ext_f.pack(fill="x", pady=4)
        ctk.CTkLabel(ext_f, text="บันทึกเป็นนามสกุล:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        for e in (".txt", ".md", ".docx"):
            ctk.CTkRadioButton(ext_f, text=e, variable=self.ext_var, value=e,
                               font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=8)

        bk_run_btn(self, "▶  เริ่มแยกตอน", self._run).pack(pady=18)

    def _browse_src(self):
        if self.mode_var.get() == "file":
            p = filedialog.askopenfilename(filetypes=[("ไฟล์นิยาย", "*.txt *.md *.docx *.pdf")])
        else:
            p = filedialog.askdirectory()
        if p: self.src_var.set(p)

    def _run(self):
        src = self.src_var.get().strip()
        out = self.out_var.get().strip()
        if not src or not out:
            messagebox.showwarning("By Keawgood", "กรุณากรอกข้อมูลให้ครบก่อนนะครับ"); return
        self.log.clear()
        threading.Thread(
            target=bk_split_to_chapters,
            args=(src, out, self.ext_var.get(), self.mode_var.get() == "folder",
                  self.prefix_var.get(), self.include_title.get(), self.log.write),
            daemon=True
        ).start()


class BK_BatchMergeTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var  = ctk.StringVar()
        self.out_var  = ctk.StringVar()
        self.batch_var = ctk.StringVar(value="5")
        self.prefix_var = ctk.StringVar(value="Chapter")
        self.ext_var  = ctk.StringVar(value=".txt")
        self._build()

    def _build(self):
        bk_section_title(self, "  📚  รวมกลุ่ม (ทีละ N ไฟล์)")
        
        bk_row(self, "โฟลเดอร์ต้นทาง:", self.src_var,
               lambda: self.src_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์ที่มีไฟล์ย่อย")))
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory(title="เลือกโฟลเดอร์บันทึก")))
        
        opt_f = ctk.CTkFrame(self, fg_color="transparent")
        opt_f.pack(fill="x", pady=4)
        
        ctk.CTkLabel(opt_f, text="จำนวนไฟล์ต่อ 1 กลุ่ม:", font=("Kanit", 13), text_color=BK_FG_MUTED, width=160, anchor="w").grid(row=0, column=0, pady=4, sticky="w")
        ctk.CTkEntry(opt_f, textvariable=self.batch_var, font=("Kanit", 12), width=80, fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8).grid(row=0, column=1, pady=4, sticky="w")
        
        ctk.CTkLabel(opt_f, text="ตั้งชื่อไฟล์ (เช่น Chapter):", font=("Kanit", 13), text_color=BK_FG_MUTED, width=160, anchor="w").grid(row=1, column=0, pady=4, sticky="w")
        ctk.CTkEntry(opt_f, textvariable=self.prefix_var, font=("Kanit", 12), width=180, fg_color=BK_BG_INPUT, border_color=BK_ACCENT, corner_radius=8).grid(row=1, column=1, pady=4, sticky="w")

        ext_f = ctk.CTkFrame(self, fg_color="transparent")
        ext_f.pack(fill="x", pady=4)
        ctk.CTkLabel(ext_f, text="บันทึกเป็นนามสกุล:", font=("Kanit", 13), text_color=BK_FG_MUTED, width=160, anchor="w").pack(side="left")
        for e in (".txt", ".md", ".docx"):
            ctk.CTkRadioButton(ext_f, text=e, variable=self.ext_var, value=e, font=("Kanit", 12), fg_color=BK_ACCENT).pack(side="left", padx=8)

        bk_run_btn(self, "▶  เริ่มรวมไฟล์เป็นกลุ่ม", self._run).pack(pady=18)

    def _run(self):
        src = self.src_var.get().strip()
        out = self.out_var.get().strip()
        if not src or not out:
            messagebox.showwarning("By Keawgood", "กรุณาเลือกโฟลเดอร์ให้ครบก่อนนะครับ")
            return
        self.log.clear()
        threading.Thread(target=bk_batch_merge_files,
                         args=(src, out, self.batch_var.get(), self.prefix_var.get(), self.ext_var.get(), self.log.write),
                         daemon=True).start()


class BK_ConvertTab(ctk.CTkFrame):
    def __init__(self, master, log):
        super().__init__(master, fg_color=BK_BG_CARD, corner_radius=14)
        self.log = log
        self.src_var  = ctk.StringVar()
        self.out_var  = ctk.StringVar()
        self.mode_var = ctk.StringVar(value="0")
        self._build()

    def _build(self):
        bk_section_title(self, "  🔄  แปลงรูปแบบไฟล์ (Convert)")
        bk_row(self, "โฟลเดอร์ต้นทาง:", self.src_var,
               lambda: self.src_var.set(filedialog.askdirectory()))
        bk_row(self, "โฟลเดอร์ปลายทาง:", self.out_var,
               lambda: self.out_var.set(filedialog.askdirectory()))
        ctk.CTkLabel(self, text="รูปแบบการแปลง:", font=("Kanit", 13),
                     text_color=BK_FG_MUTED).pack(anchor="w", pady=(10, 4))
        grid = ctk.CTkFrame(self, fg_color="transparent")
        grid.pack(fill="x", padx=4)
        for i, (_, _, label) in enumerate(BK_CONVERSIONS):
            ctk.CTkRadioButton(grid, text=label, variable=self.mode_var, value=str(i),
                               font=("Kanit", 13, "bold"), fg_color=BK_ACCENT,
                               text_color=BK_FG_TEXT).grid(row=i//2, column=i%2, padx=20, pady=6, sticky="w")
        bk_run_btn(self, "▶  เริ่มแปลงไฟล์", self._run).pack(pady=18)

    def _run(self):
        src = self.src_var.get().strip()
        out = self.out_var.get().strip()
        idx = int(self.mode_var.get())
        from_ext, to_ext, _ = BK_CONVERSIONS[idx]
        if not src or not out:
            messagebox.showwarning("By Keawgood", "กรุณากรอกข้อมูลให้ครบก่อนนะครับ"); return
        self.log.clear()
        threading.Thread(target=bk_convert_files,
                         args=(src, out, from_ext, to_ext, self.log.write),
                         daemon=True).start()


class ByKeawgoodWindow(ctk.CTkToplevel):
    def __init__(self, master):
        super().__init__(master)
        self.title("By Keawgood — Thai Novel Manager")
        self.geometry("940x760")
        self.minsize(820, 640)
        self.configure(fg_color=BK_BG_DARK)
        self._build()
        self.lift()
        self.focus_force()

    def _build(self):
        # ── Gradient accent strip ──────────────────────────────────────────────
        accent_strip = ctk.CTkFrame(self, height=4, corner_radius=0,
                                    fg_color=BK_ACCENT)
        accent_strip.pack(fill="x")

        hdr = ctk.CTkFrame(self, fg_color=BK_BG_CARD, corner_radius=0, height=72)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        # Icon badge
        badge = ctk.CTkFrame(hdr, fg_color=("#E0E7FF","#1E1B4B"), corner_radius=10,
                              width=46, height=46)
        badge.pack(side="left", padx=(20, 10), pady=13)
        badge.pack_propagate(False)
        ctk.CTkLabel(badge, text="⧉", font=("Tahoma", 22, "bold"),
                     text_color=BK_ACCENT2).place(relx=0.5, rely=0.5, anchor="center")

        title_col = ctk.CTkFrame(hdr, fg_color="transparent")
        title_col.pack(side="left")
        ctk.CTkLabel(title_col, text="By Keawgood", font=("Kanit", 20, "bold"),
                     text_color=BK_ACCENT2).pack(anchor="w")
        ctk.CTkLabel(title_col, text="Thai Novel File Manager",
                     font=("Kanit", 11), text_color=BK_FG_MUTED).pack(anchor="w")

        self._mode = ctk.StringVar(value="dark")
        self.mode_switch = ctk.CTkSwitch(hdr, text="🌙 Dark", variable=self._mode,
                                         onvalue="dark", offvalue="light",
                                         command=self._toggle_mode,
                                         font=("Kanit", 12), fg_color=BK_ACCENT,
                                         progress_color=BK_ACCENT2)
        self.mode_switch.pack(side="right", padx=24)

        body = ctk.CTkFrame(self, fg_color="transparent")
        body.pack(fill="both", expand=True, padx=18, pady=14)
        body.columnconfigure(0, weight=1)
        body.rowconfigure(0, weight=1)

        tabview = ctk.CTkTabview(body, fg_color=BK_BG_CARD,
                                 segmented_button_fg_color=BK_BG_INPUT,
                                 segmented_button_selected_color=BK_ACCENT,
                                 segmented_button_selected_hover_color=BK_ACCENT2,
                                 segmented_button_unselected_color=BK_BG_INPUT,
                                 text_color=BK_FG_TEXT, corner_radius=14)
        tabview.grid(row=0, column=0, sticky="nsew", pady=(0, 10))

        tab_names = ("🗂️  รวมไฟล์ทั้งหมด", "📚  รวมกลุ่ม (N ไฟล์)", "✂️  แยกตอน", "🔄  แปลงไฟล์")
        for name in tab_names:
            tabview.add(name)
            tabview.tab(name).configure(fg_color=BK_BG_CARD)

        # ── Log section ───────────────────────────────────────────────────────
        log_frame = ctk.CTkFrame(body, fg_color=BK_BG_CARD, corner_radius=12)
        log_frame.grid(row=1, column=0, sticky="nsew")
        body.rowconfigure(1, weight=0, minsize=200)

        log_hdr = ctk.CTkFrame(log_frame, fg_color="transparent", height=36)
        log_hdr.pack(fill="x", padx=14, pady=(8, 0))
        log_hdr.pack_propagate(False)
        log_hdr.columnconfigure(0, weight=1)

        ctk.CTkLabel(log_hdr, text="📋  Console Log",
                     font=("Kanit", 13, "bold"), text_color=BK_ACCENT2,
                     anchor="w").pack(side="left")

        ctk.CTkButton(log_hdr, text="🗑  ล้าง Log", width=110,
                      fg_color=("#CBD5E1", "#3F3F5A"), hover_color=BK_ACCENT,
                      font=("Kanit", 11), corner_radius=8,
                      command=self.log.clear if hasattr(self, 'log') else lambda: None
                      ).pack(side="right")

        self.log = BK_LogBox(log_frame, height=160)
        self.log.pack(fill="both", expand=True, padx=10, pady=(4, 10))

        # Re-wire clear button now that self.log exists
        for w in log_hdr.winfo_children():
            if isinstance(w, ctk.CTkButton):
                w.configure(command=self.log.clear)

        cls_tab_map = [
            (BK_MergeTab,      "🗂️  รวมไฟล์ทั้งหมด"),
            (BK_BatchMergeTab, "📚  รวมกลุ่ม (N ไฟล์)"),
            (BK_SplitTab,      "✂️  แยกตอน"),
            (BK_ConvertTab,    "🔄  แปลงไฟล์"),
        ]
        for cls, tab_name in cls_tab_map:
            widget = cls(tabview.tab(tab_name), self.log)
            widget.pack(fill="both", expand=True, padx=16, pady=12)

        self.log.write("👋  ยินดีต้อนรับสู่ By Keawgood  v5.0!")
        self.log.write("   เลือก Tab ที่ต้องการแล้วกด ▶ เริ่ม\n")

    def _toggle_mode(self):
        mode = self._mode.get()
        ctk.set_appearance_mode(mode)
        self.mode_switch.configure(text="☀️  Light" if mode == "light" else "🌙  Dark")



# ══════════════════════════════════════════════════════════════════════════════
#  GLOBAL THEME SYSTEM + DnD BOOTSTRAP
#  Themes: DARK · LIGHT · SEPIA · AMOLED · NORD
# ══════════════════════════════════════════════════════════════════════════════

# ── DnD Import ────────────────────────────────────────────────────────────────
try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    DND_OK = True
except ImportError:
    DND_OK = False

# ── Natural sort helper (fixes audio file ordering) ──────────────────────────
def _natural_key(path: str):
    """Sort paths naturally: 'ep2' < 'ep10', handles Thai chars."""
    name = os.path.basename(path)
    parts = re.split(r'(\d+)', name)
    return [int(p) if p.isdigit() else p.lower() for p in parts]

# ── Theme definitions ─────────────────────────────────────────────────────────
#  Each theme is a flat dict used by ALL sub-windows
THEMES = {
    "DARK": {
        "label":      "🌑  Dark",
        "ctk_mode":   "dark",
        "bg_root":    "#0A0A0F",
        "bg_sidebar": "#0F0F17",
        "bg_card":    "#13131E",
        "bg_card_hv": "#1A1A2B",
        "bg_input":   "#1E1E30",
        "bg_border":  "#252542",
        "fg_primary": "#F0F0FF",
        "fg_body":    "#B0B0D0",
        "fg_muted":   "#50507A",
        "accent":     "#6366F1",
        "accent2":    "#818CF8",
        "success":    "#22D3A0",
        "warn":       "#FBBF24",
        "error":      "#F87171",
        "console_bg": "#0D0D1A",
        "console_fg": "#A0FFB0",
    },
    "LIGHT": {
        "label":      "☀️  Light",
        "ctk_mode":   "light",
        "bg_root":    "#F5F7FF",
        "bg_sidebar": "#ECEEF8",
        "bg_card":    "#FFFFFF",
        "bg_card_hv": "#F0F2FF",
        "bg_input":   "#E8EBF8",
        "bg_border":  "#D0D4F0",
        "fg_primary": "#0F1030",
        "fg_body":    "#3A3D60",
        "fg_muted":   "#8088AA",
        "accent":     "#4F46E5",
        "accent2":    "#6366F1",
        "success":    "#059669",
        "warn":       "#D97706",
        "error":      "#DC2626",
        "console_bg": "#F0F4FF",
        "console_fg": "#1E293B",
    },
    "SEPIA": {
        "label":      "🌿  Eye Care",
        "ctk_mode":   "light",
        "bg_root":    "#1A1610",
        "bg_sidebar": "#211D15",
        "bg_card":    "#27221A",
        "bg_card_hv": "#322B20",
        "bg_input":   "#302820",
        "bg_border":  "#403830",
        "fg_primary": "#F0E6C8",
        "fg_body":    "#C8B890",
        "fg_muted":   "#7A6848",
        "accent":     "#C8963C",
        "accent2":    "#E0B060",
        "success":    "#6BAA60",
        "warn":       "#D4A030",
        "error":      "#C06060",
        "console_bg": "#1A1510",
        "console_fg": "#D0C090",
    },
    "AMOLED": {
        "label":      "⚫  AMOLED",
        "ctk_mode":   "dark",
        "bg_root":    "#000000",
        "bg_sidebar": "#050505",
        "bg_card":    "#0A0A0A",
        "bg_card_hv": "#111111",
        "bg_input":   "#0D0D0D",
        "bg_border":  "#1A1A1A",
        "fg_primary": "#FFFFFF",
        "fg_body":    "#CCCCCC",
        "fg_muted":   "#555555",
        "accent":     "#7C6BFF",
        "accent2":    "#9F8FFF",
        "success":    "#00E676",
        "warn":       "#FFD740",
        "error":      "#FF5252",
        "console_bg": "#000000",
        "console_fg": "#00E676",
    },
    "NORD": {
        "label":      "❄️  Nord",
        "ctk_mode":   "dark",
        "bg_root":    "#2E3440",
        "bg_sidebar": "#252B37",
        "bg_card":    "#3B4252",
        "bg_card_hv": "#434C5E",
        "bg_input":   "#2E3440",
        "bg_border":  "#4C566A",
        "fg_primary": "#ECEFF4",
        "fg_body":    "#D8DEE9",
        "fg_muted":   "#4C566A",
        "accent":     "#88C0D0",
        "accent2":    "#81A1C1",
        "success":    "#A3BE8C",
        "warn":       "#EBCB8B",
        "error":      "#BF616A",
        "console_bg": "#242933",
        "console_fg": "#A3BE8C",
    },
}
THEME_KEYS = list(THEMES.keys())

# Global active theme key — all windows read this
_ACTIVE_THEME_KEY = "DARK"
_THEME_CALLBACKS: list = []   # list of callables to notify on theme change

def get_theme() -> dict:
    return THEMES[_ACTIVE_THEME_KEY]

def set_theme(key: str):
    global _ACTIVE_THEME_KEY
    _ACTIVE_THEME_KEY = key
    T = get_theme()
    ctk.set_appearance_mode(T["ctk_mode"])
    for cb in _THEME_CALLBACKS:
        try:
            cb(key)
        except Exception:
            pass

def register_theme_callback(cb):
    _THEME_CALLBACKS.append(cb)

def unregister_theme_callback(cb):
    try:
        _THEME_CALLBACKS.remove(cb)
    except ValueError:
        pass

# ── Language system ────────────────────────────────────────────────────────────
_ACTIVE_LANG = "TH"
_LANG_CALLBACKS: list = []

def get_lang() -> str:
    return _ACTIVE_LANG

def set_lang(lang: str):
    global _ACTIVE_LANG
    _ACTIVE_LANG = lang
    for cb in _LANG_CALLBACKS:
        try: cb(lang)
        except Exception: pass

def register_lang_callback(cb):
    _LANG_CALLBACKS.append(cb)

def unregister_lang_callback(cb):
    try: _LANG_CALLBACKS.remove(cb)
    except ValueError: pass

# ── Shared UI helpers ──────────────────────────────────────────────────────────
FONT_TH   = "Tahoma"
FONT_MONO = "Consolas"

def mk_font(size=14, bold=False, mono=False):
    fam = FONT_MONO if mono else FONT_TH
    wt  = "bold" if bold else "normal"
    return ctk.CTkFont(family=fam, size=size, weight=wt)

def tk_font(size=13, bold=False, mono=False):
    fam = FONT_MONO if mono else FONT_TH
    wt  = "bold" if bold else "normal"
    return (fam, size, wt)

# ── DnD utility: patch a plain tk widget to accept drops ──────────────────────
def _dnd_register(widget, callback):
    """
    Safely register a plain tk.Widget for DnD drops.
    Works only if DND_OK and the Tcl DnD extension is loaded on the root.
    """
    if not DND_OK:
        return
    try:
        widget.tk.call('tkdnd::drop_target', 'register', widget._w, '{DND_Files}')
        widget.bind('<<Drop:DND_Files>>', callback)
        widget.bind('<<Drop>>', callback)
    except Exception:
        pass

def _dnd_require_root(tk_root):
    """Load tkdnd on the Tk root — must be called once after mainloop starts."""
    if not DND_OK:
        return False
    try:
        TkinterDnD._require(tk_root)
        return True
    except Exception as e:
        print(f"DnD _require failed: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  VOCAB OPTIMIZER WINDOW  — v3
# ══════════════════════════════════════════════════════════════════════════════

class VocabOptimizerWindow(ctk.CTkToplevel):
    """
    จัดการคำศัพท์นิยาย
    • Drag & Drop ไฟล์ .txt (ต้องการ tkinterdnd2)
    • Multi-theme aware
    • ปุ่มเปิดไฟล์หลายไฟล์พร้อมกัน
    """

    _KEEP_OPTS = [
        "เก็บอันที่ยาวที่สุด (แนะนำ)",
        "เก็บอันล่าสุดที่เจอ",
        "เก็บอันแรกที่เจอ",
    ]

    def __init__(self, master):
        super().__init__(master)
        self.title("✦ Vocab Optimizer  v5.0  ·  By Keawgood")
        self.geometry("1260x800")
        self.minsize(900, 600)
        self.font_size = 18
        self._dnd_active = False
        self._T = get_theme()

        register_theme_callback(self._on_theme_change)
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        self._build()
        self._apply_theme()

        # DnD — register after widget is mapped
        self.after(300, self._init_dnd)
        self.lift(); self.focus_force()

    # ── Close ──────────────────────────────────────────────────────────────────
    def _on_close(self):
        unregister_theme_callback(self._on_theme_change)
        self.destroy()

    def _on_theme_change(self, key):
        self._T = THEMES[key]
        self.after(50, self._apply_theme)

    # ── Build ──────────────────────────────────────────────────────────────────
    def _build(self):
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        self._build_sidebar()
        self._build_editor()

    # ── Sidebar ────────────────────────────────────────────────────────────────
    def _build_sidebar(self):
        T = self._T
        self._sb = ctk.CTkFrame(self, width=248, corner_radius=0)
        self._sb.grid(row=0, column=0, sticky="nsew")
        self._sb.grid_propagate(False)
        self._sb.grid_rowconfigure(15, weight=1)

        # Logo strip
        self._logo = ctk.CTkFrame(self._sb, height=80, corner_radius=0)
        self._logo.grid(row=0, column=0, sticky="ew")
        self._logo.grid_propagate(False)
        self._logo_title = ctk.CTkLabel(
            self._logo, text="✦  Vocab Optimizer",
            font=mk_font(19, bold=True))
        self._logo_title.place(relx=0.5, rely=0.42, anchor="center")
        self._logo_sub = ctk.CTkLabel(
            self._logo, text="Novel Toolkit · By Keawgood",
            font=mk_font(11))
        self._logo_sub.place(relx=0.5, rely=0.74, anchor="center")

        # Buttons
        self._btn_open = self._sb_btn("📂  เปิดไฟล์ .txt", self.open_file, row=1)
        self._btn_save = self._sb_btn("💾  บันทึกผลลัพธ์", self.save_file, row=2)
        self._btn_clear = self._sb_btn("🗑  เคลียร์ทั้งหมด", self.clear_text, row=3,
                                        color_key="error")

        self._div(row=4)

        # Keep mode
        ctk.CTkLabel(self._sb, text="⚙  การจัดการคำซ้ำ",
                     font=mk_font(12, bold=True)
                     ).grid(row=5, column=0, padx=18, pady=(2, 4), sticky="w")
        self._keep_var = ctk.StringVar(value=self._KEEP_OPTS[0])
        self._keep_menu = ctk.CTkOptionMenu(
            self._sb, variable=self._keep_var,
            values=self._KEEP_OPTS,
            font=mk_font(13), width=220, height=34, corner_radius=8)
        self._keep_menu.grid(row=6, column=0, padx=16, pady=(0, 10), sticky="ew")

        self._div(row=7)

        # Font size
        ctk.CTkLabel(self._sb, text="🔤  ขนาดตัวอักษร",
                     font=mk_font(12, bold=True)
                     ).grid(row=8, column=0, padx=18, pady=(2, 4), sticky="w")
        fr = ctk.CTkFrame(self._sb, fg_color="transparent")
        fr.grid(row=9, column=0, padx=16, pady=(0, 10), sticky="ew")
        fr.columnconfigure((0, 1, 2), weight=1)
        self._btn_fa = ctk.CTkButton(fr, text="A−", width=56, height=30, corner_radius=8,
                                      command=self.dec_font)
        self._btn_fa.grid(row=0, column=0, padx=(0, 3))
        self._lbl_fsize = ctk.CTkLabel(fr, text=f"{self.font_size}pt",
                                        font=mk_font(12, bold=True))
        self._lbl_fsize.grid(row=0, column=1)
        self._btn_fb = ctk.CTkButton(fr, text="A+", width=56, height=30, corner_radius=8,
                                      command=self.inc_font)
        self._btn_fb.grid(row=0, column=2, padx=(3, 0))

        self._div(row=10)

        # Process button
        self._btn_proc = ctk.CTkButton(
            self._sb, text="⚡  ตัดคำซ้ำ  →",
            font=mk_font(18, bold=True), height=52, corner_radius=12,
            command=self.process_text)
        self._btn_proc.grid(row=16, column=0, padx=16, pady=(0, 16), sticky="sew")

    def _sb_btn(self, text, cmd, row, color_key="accent"):
        b = ctk.CTkButton(self._sb, text=text, font=mk_font(14, bold=True),
                          height=40, corner_radius=10, command=cmd)
        b.grid(row=row, column=0, padx=16, pady=5, sticky="ew")
        return b

    def _div(self, row):
        ctk.CTkFrame(self._sb, height=1, corner_radius=0
                     ).grid(row=row, column=0, sticky="ew", padx=14, pady=8)

    # ── Editor area ────────────────────────────────────────────────────────────
    def _build_editor(self):
        self._main = ctk.CTkFrame(self, fg_color="transparent")
        self._main.grid(row=0, column=1, sticky="nsew", padx=(0, 14), pady=14)
        self._main.grid_columnconfigure((0, 1), weight=1)
        self._main.grid_rowconfigure(1, weight=1)

        # Column headers
        for col, icon, title in [(0, "📥", "ต้นฉบับ  —  ลากไฟล์ .txt มาวางที่นี่"),
                                  (1, "📤", "ผลลัพธ์หลังตัดคำซ้ำ")]:
            h = ctk.CTkFrame(self._main, height=38, corner_radius=8)
            h.grid(row=0, column=col,
                   padx=(0 if col else 0, 6 if col == 0 else 0),
                   pady=(0, 5), sticky="ew")
            h.grid_propagate(False)
            ctk.CTkLabel(h, text=f"{icon}  {title}",
                         font=mk_font(13, bold=True)
                         ).place(relx=0.02, rely=0.5, anchor="w")
            self._main.grid_columnconfigure(col, weight=1)

        # Input textbox
        self._in_frame = ctk.CTkFrame(self._main, corner_radius=12)
        self._in_frame.grid(row=1, column=0, padx=(0, 6), sticky="nsew")
        self._in_frame.grid_rowconfigure(0, weight=1)
        self._in_frame.grid_columnconfigure(0, weight=1)

        self.txt_input = ctk.CTkTextbox(
            self._in_frame, font=mk_font(self.font_size),
            wrap="none", corner_radius=10,
            border_width=2)
        self.txt_input.grid(sticky="nsew", padx=2, pady=2)
        self.txt_input.insert("1.0",
            "╔══════════════════════════════════════════╗\n"
            "║  📂  ลากไฟล์ .txt มาวางที่นี่             ║\n"
            "║      หรือกดปุ่ม 'เปิดไฟล์' ทางซ้าย       ║\n"
            "║                                          ║\n"
            "║  ▸  รองรับหลายไฟล์พร้อมกัน              ║\n"
            "║  ▸  กด ⚡ เพื่อตัดคำซ้ำทันที            ║\n"
            "╚══════════════════════════════════════════╝")

        # Output textbox
        self.txt_output = ctk.CTkTextbox(
            self._main, font=mk_font(self.font_size),
            wrap="none", corner_radius=12,
            border_width=2)
        self.txt_output.grid(row=1, column=1, sticky="nsew")
        self.txt_output.insert("1.0", "— ผลลัพธ์จะแสดงที่นี่ —")

        # Status bar
        self._status_var = ctk.StringVar(value="พร้อมใช้งาน")
        self._statusbar = ctk.CTkFrame(self._main, height=30, corner_radius=6)
        self._statusbar.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(6, 0))
        self._statusbar.grid_propagate(False)
        self._status_lbl = ctk.CTkLabel(
            self._statusbar, textvariable=self._status_var,
            font=mk_font(12, mono=True), anchor="w")
        self._status_lbl.place(relx=0.01, rely=0.5, anchor="w")

    # ── DnD init ───────────────────────────────────────────────────────────────
    def _init_dnd(self):
        """
        Correct DnD setup for tkinterdnd2 with CTk:
        The Tk root must have TkinterDnD loaded, then each target widget
        is registered using the Tcl tkdnd API directly.
        """
        if not DND_OK:
            self._status_var.set("⚠  DnD: tkinterdnd2 ไม่ได้ติดตั้ง  →  pip install tkinterdnd2")
            return
        try:
            # Register the Toplevel itself
            self.drop_target_register(DND_FILES)
            self.dnd_bind('<<Drop>>', self._on_drop)
            self.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self.dnd_bind('<<DragLeave>>', self._on_drag_leave)

            # Also register the inner _textbox widget
            inner = self.txt_input._textbox
            inner.drop_target_register(DND_FILES)
            inner.dnd_bind('<<Drop>>', self._on_drop)

            self._status_var.set("✅  พร้อมใช้งาน — ลากไฟล์ .txt มาวางได้เลย")
        except Exception as e:
            # Fallback: try raw Tcl API
            try:
                for w in [self, self.txt_input._textbox]:
                    self.tk.call('tkdnd::drop_target', 'register', w, 'DND_Files')
                    w.bind('<<Drop>>', self._on_drop)
                self._status_var.set("✅  DnD พร้อม (Tcl fallback)")
            except Exception as e2:
                self._status_var.set(f"⚠  DnD init error: {e2}")

    def _on_drag_enter(self, event=None):
        if not self._dnd_active:
            self._dnd_active = True
            T = self._T
            self.txt_input.configure(border_color=T["accent"], border_width=3)
            self._status_var.set("🎯  วางไฟล์ .txt ได้เลย!")

    def _on_drag_leave(self, event=None):
        self._dnd_active = False
        self._apply_theme()
        self._status_var.set("พร้อมใช้งาน — ลากไฟล์ .txt มาวางได้เลย")

    def _on_drop(self, event):
        self._dnd_active = False
        self._apply_theme()
        try:
            raw = event.data if hasattr(event, 'data') else event.widget.tk.splitlist(str(event))
            files = self.tk.splitlist(raw) if isinstance(raw, str) else list(raw)
        except Exception:
            return
        txt_files = [f.strip("{}") for f in files if f.strip("{}").lower().endswith(".txt")]
        if not txt_files:
            messagebox.showwarning("Vocab Optimizer", "กรุณาลากเฉพาะไฟล์ .txt เท่านั้น")
            return
        self._load_files(txt_files)

    # ── File I/O ───────────────────────────────────────────────────────────────
    def _read_txt(self, path):
        for enc in ("utf-8-sig", "utf-8", "tis-620", "cp874"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        raise ValueError(f"อ่านไฟล์ไม่ได้: {path}")

    def _load_files(self, paths):
        parts = []
        for fp in paths:
            try:
                parts.append(self._read_txt(fp))
            except Exception as e:
                messagebox.showerror("Error", str(e))
                return
        text = "\n".join(parts)
        self.txt_input.delete("1.0", "end")
        self.txt_input.insert("end", text)
        self._status_var.set(
            f"✅  โหลด {len(paths)} ไฟล์  ·  {len(text):,} ตัวอักษร")

    def open_file(self):
        fps = filedialog.askopenfilenames(
            title="เลือกไฟล์คำศัพท์",
            filetypes=[("Text", "*.txt"), ("All", "*.*")])
        if fps:
            self._load_files(list(fps))

    def save_file(self):
        fp = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text", "*.txt")])
        if fp:
            try:
                with open(fp, "w", encoding="utf-8-sig") as f:
                    f.write(self.txt_output.get("1.0", "end"))
                self._status_var.set(f"💾  บันทึกแล้ว → {os.path.basename(fp)}")
            except Exception as e:
                messagebox.showerror("Error", str(e))

    def clear_text(self):
        self.txt_input.delete("1.0", "end")
        self.txt_output.delete("1.0", "end")
        self._status_var.set("🗑  ล้างแล้ว")

    # ── Font ───────────────────────────────────────────────────────────────────
    def inc_font(self):
        self.font_size = min(36, self.font_size + 2)
        self._refresh_font()

    def dec_font(self):
        self.font_size = max(10, self.font_size - 2)
        self._refresh_font()

    def _refresh_font(self):
        f = mk_font(self.font_size)
        self.txt_input.configure(font=f)
        self.txt_output.configure(font=f)
        self._lbl_fsize.configure(text=f"{self.font_size}pt")

    # ── Process ────────────────────────────────────────────────────────────────
    def process_text(self):
        text = self.txt_input.get("1.0", "end")
        lines = text.split("\n")
        vocab: dict = {}
        mode = self._keep_var.get()
        orig = 0
        for i, line in enumerate(lines):
            orig_line = line.strip()
            if not orig_line or orig_line.startswith("---") or orig_line[0] in "╔║╚":
                continue
            orig += 1
            if orig_line.startswith("[") and orig_line.endswith("]"):
                vocab[f"__H_{i}"] = {"line": orig_line, "order": i, "h": True}
                continue
            clean = re.sub(r"^[\-\s]+", "", orig_line)
            parts = re.split(r"[\s=/]+", clean, maxsplit=1)
            key = parts[0].strip() if parts else ""
            if not key:
                continue
            if key in vocab and not vocab[key].get("h"):
                if mode == self._KEEP_OPTS[0] and len(orig_line) > len(vocab[key]["line"]):
                    vocab[key]["line"] = orig_line
                elif mode == self._KEEP_OPTS[1]:
                    vocab[key]["line"] = orig_line
            else:
                vocab[key] = {"line": orig_line, "order": i, "h": False}

        sv = sorted(vocab.values(), key=lambda x: x["order"])
        result = []
        for v in sv:
            if v.get("h") and result and result[-1]:
                result.append("")
            result.append(v["line"])

        after = sum(1 for v in sv if not v.get("h"))
        self.txt_output.delete("1.0", "end")
        self.txt_output.insert("end", "\n".join(result))
        self._status_var.set(
            f"⚡  เสร็จ! ก่อน {orig:,} → หลัง {after:,} บรรทัด  (ลบ {orig - after:,} รายการซ้ำ)")

    # ── Theme apply ────────────────────────────────────────────────────────────
    def _apply_theme(self):
        T = self._T
        self.configure(fg_color=T["bg_root"])
        self._sb.configure(fg_color=T["bg_sidebar"])
        self._logo.configure(fg_color=T["bg_sidebar"])
        self._logo_title.configure(text_color=T["accent"])
        self._logo_sub.configure(text_color=T["fg_muted"])
        self._in_frame.configure(fg_color=T["bg_card"])
        self.txt_input.configure(
            fg_color=T["bg_card"],
            text_color=T["fg_body"],
            border_color=T["bg_border"])
        self.txt_output.configure(
            fg_color=T["bg_card"],
            text_color=T["fg_body"],
            border_color=T["bg_border"])
        self._statusbar.configure(fg_color=T["bg_input"])
        self._status_lbl.configure(text_color=T["fg_muted"])
        for btn in [self._btn_open, self._btn_save, self._btn_proc, self._btn_fa, self._btn_fb]:
            btn.configure(fg_color=T["accent"], hover_color=T["accent2"],
                          text_color=T["fg_primary"])
        self._btn_clear.configure(fg_color=T["error"], hover_color="#B91C1C",
                                   text_color="#FFFFFF")
        self._keep_menu.configure(
            fg_color=T["bg_input"],
            button_color=T["accent"],
            button_hover_color=T["accent2"],
            text_color=T["fg_body"])


# ══════════════════════════════════════════════════════════════════════════════
#  AUDIO BY KEAWGOOD WINDOW  — v3
#  • Natural sort (เรียงลำดับตัวเลขถูกต้อง: 1,2,...10 ไม่ใช่ 1,10,2)
#  • Drag & Drop หลายไฟล์ m4a/mp3/wav + PNG/JPG พร้อมกัน
#  • เลือกรูปภาพพื้นหลัง 1 รูป หรือลากมาวางได้
#  • Reorder ลากเรียงลำดับใน listbox ได้
#  • Multi-theme aware
# ══════════════════════════════════════════════════════════════════════════════

_AUDIO_I18N = {
    "TH": {
        "title":       "🎵  Audio → Video  ·  By Keawgood",
        "sub":         "แปลงไฟล์เสียงเป็นวิดีโอ (Batch)",
        "sec_audio":   "🎵  ไฟล์เสียง",
        "sec_image":   "🖼  รูปภาพพื้นหลัง",
        "sec_output":  "📁  โฟลเดอร์บันทึก",
        "sec_settings":"⚙  ตั้งค่า",
        "sec_progress":"📊  ความคืบหน้า",
        "btn_add":     "➕  เพิ่มไฟล์เสียง",
        "btn_clear":   "🗑  ล้างรายการ",
        "btn_sort_az":  "🔤  A→Z",
        "btn_sort_nat": "🔢  123",
        "btn_img":     "เลือกรูปภาพ",
        "btn_out":     "เลือกโฟลเดอร์",
        "btn_start":   "⚡  สร้างวิดีโอ",
        "lbl_fname":   "ชื่อไฟล์ผลลัพธ์:",
        "lbl_chunk":   "คลิปต่อ 1 ไฟล์วิดีโอ:",
        "hint_drag":   "ลากไฟล์ .m4a/.mp3/.wav หรือรูปภาพมาวางที่นี่",
        "no_img":      "ยังไม่ได้เลือกรูป",
        "no_out":      "ยังไม่ได้เลือกโฟลเดอร์",
        "err_missing": "กรุณาเลือกไฟล์เสียง รูปภาพ และโฟลเดอร์ให้ครบก่อนนะครับ",
        "err_num":     "จำนวนคลิปต้องเป็นตัวเลขที่มากกว่า 0",
        "err_moviepy": "กรุณาติดตั้ง moviepy ก่อน:\npip install moviepy",
        "status_idle": "พร้อมทำงาน",
        "status_done": "✅  เสร็จสมบูรณ์!",
        "status_proc": "กำลังเรนเดอร์ {c}/{t} ...",
        "cnt_files":   "{n} ไฟล์เสียง",
        "lang_switch": "EN",
    },
    "EN": {
        "title":       "🎵  Audio → Video  ·  By Keawgood",
        "sub":         "Batch Audio-to-Video Converter",
        "sec_audio":   "🎵  Audio Files",
        "sec_image":   "🖼  Background Image",
        "sec_output":  "📁  Output Folder",
        "sec_settings":"⚙  Settings",
        "sec_progress":"📊  Progress",
        "btn_add":     "➕  Add Audio",
        "btn_clear":   "🗑  Clear",
        "btn_sort_az":  "🔤  A→Z",
        "btn_sort_nat": "🔢  1-2-3",
        "btn_img":     "Select Image",
        "btn_out":     "Select Folder",
        "btn_start":   "⚡  Create Video",
        "lbl_fname":   "Output filename:",
        "lbl_chunk":   "Clips per video file:",
        "hint_drag":   "Drag .m4a/.mp3/.wav or image files here",
        "no_img":      "No image selected",
        "no_out":      "No output folder",
        "err_missing": "Please select audio files, image, and output folder.",
        "err_num":     "Clip count must be a positive integer.",
        "err_moviepy": "Please install moviepy first:\npip install moviepy",
        "status_idle": "Ready",
        "status_done": "✅  Done!",
        "status_proc": "Rendering {c}/{t} ...",
        "cnt_files":   "{n} audio files",
        "lang_switch": "ภาษาไทย",
    },
}


class AudioByKeawgoodWindow(ctk.CTkToplevel):
    """
    Batch Audio → Video Converter
    UI: 2-column layout — file list (left) + controls (right)
    """

    def __init__(self, master):
        super().__init__(master)
        self._lang      = get_lang()
        self.audio_paths: list = []
        self.image_path = ""
        self.output_dir = ""
        self._T         = get_theme()
        self._processing = False

        register_theme_callback(self._on_theme_change)
        register_lang_callback(self._on_lang_change)
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        self._setup_window()
        self._build_ui()
        self._apply_theme()
        self.after(300, self._init_dnd)
        self.lift(); self.focus_force()

    # ── Lifecycle ──────────────────────────────────────────────────────────────
    def _on_close(self):
        unregister_theme_callback(self._on_theme_change)
        unregister_lang_callback(self._on_lang_change)
        self.destroy()

    def _on_theme_change(self, key):
        self._T = THEMES[key]
        self.after(50, self._apply_theme)

    def _on_lang_change(self, lang):
        self._lang = lang
        self.after(50, self._refresh_lang)

    def _L(self) -> dict:
        return _AUDIO_I18N.get(self._lang, _AUDIO_I18N["TH"])

    def _setup_window(self):
        L = self._L()
        self.title(L["title"])
        self.geometry("1100x720")
        self.minsize(900, 580)

    # ── Build UI ───────────────────────────────────────────────────────────────
    def _build_ui(self):
        self.grid_columnconfigure(0, weight=3)
        self.grid_columnconfigure(1, weight=2)
        self.grid_rowconfigure(0, weight=1)

        self._build_left_panel()
        self._build_right_panel()
        self._build_topbar()

    def _build_topbar(self):
        """Overlay top bar with title + lang switch."""
        L = self._L()
        self._topbar = ctk.CTkFrame(self, height=60, corner_radius=0)
        self._topbar.grid(row=0, column=0, columnspan=2, sticky="new")
        self._topbar.grid_propagate(False)
        self._topbar.columnconfigure(1, weight=1)

        self._lbl_title = ctk.CTkLabel(
            self._topbar, text=L["title"],
            font=mk_font(22, bold=True))
        self._lbl_title.grid(row=0, column=0, padx=20, pady=10, sticky="w")

        self._lbl_sub = ctk.CTkLabel(
            self._topbar, text=L["sub"],
            font=mk_font(12))
        self._lbl_sub.grid(row=0, column=1, padx=4, pady=10, sticky="w")

        self._btn_lang = ctk.CTkButton(
            self._topbar, text=L["lang_switch"],
            font=mk_font(12, bold=True),
            width=90, height=30, corner_radius=8,
            command=self._toggle_lang)
        self._btn_lang.grid(row=0, column=2, padx=(0, 8))

    # ── Left panel: audio list ─────────────────────────────────────────────────
    def _build_left_panel(self):
        L = self._L()
        self._left = ctk.CTkFrame(self, corner_radius=0)
        self._left.grid(row=0, column=0, sticky="nsew", padx=(10, 5), pady=(70, 10))
        self._left.grid_rowconfigure(2, weight=1)
        self._left.grid_columnconfigure(0, weight=1)

        # Section header
        self._sec_audio = ctk.CTkLabel(
            self._left, text=L["sec_audio"],
            font=mk_font(15, bold=True), anchor="w")
        self._sec_audio.grid(row=0, column=0, sticky="w", padx=14, pady=(12, 6))

        # Toolbar row
        self._audio_toolbar = ctk.CTkFrame(self._left, fg_color="transparent")
        self._audio_toolbar.grid(row=1, column=0, sticky="ew", padx=10, pady=(0, 6))
        for col in range(6):
            self._audio_toolbar.columnconfigure(col, weight=1)

        self._btn_add = ctk.CTkButton(
            self._audio_toolbar, text=L["btn_add"],
            font=mk_font(12, bold=True), height=32, corner_radius=8,
            command=self._select_audio)
        self._btn_add.grid(row=0, column=0, padx=(0, 4), sticky="ew")

        self._btn_sort_nat = ctk.CTkButton(
            self._audio_toolbar, text=L["btn_sort_nat"],
            font=mk_font(11, bold=True), height=32, corner_radius=8,
            width=56,
            command=self._sort_natural)
        self._btn_sort_nat.grid(row=0, column=1, padx=2, sticky="ew")

        self._btn_sort_az = ctk.CTkButton(
            self._audio_toolbar, text=L["btn_sort_az"],
            font=mk_font(11, bold=True), height=32, corner_radius=8,
            width=56,
            command=self._sort_alpha)
        self._btn_sort_az.grid(row=0, column=2, padx=2, sticky="ew")

        self._btn_del_sel = ctk.CTkButton(
            self._audio_toolbar, text="✕ Del",
            font=mk_font(11, bold=True), height=32, corner_radius=8,
            width=56,
            command=self._delete_selected)
        self._btn_del_sel.grid(row=0, column=3, padx=2, sticky="ew")

        self._btn_clear = ctk.CTkButton(
            self._audio_toolbar, text=L["btn_clear"],
            font=mk_font(11, bold=True), height=32, corner_radius=8,
            command=self._clear_audio)
        self._btn_clear.grid(row=0, column=4, padx=(2, 0), sticky="ew")

        # Listbox container
        self._list_container = tk.Frame(self._left, bg="#1E1E30",
                                         highlightthickness=1,
                                         highlightbackground="#252542")
        self._list_container.grid(row=2, column=0, sticky="nsew", padx=10, pady=(0, 6))
        self._list_container.grid_rowconfigure(0, weight=1)
        self._list_container.grid_columnconfigure(0, weight=1)

        self._audio_lb = tk.Listbox(
            self._list_container,
            selectmode=tk.EXTENDED,
            font=tk_font(13, mono=True),
            relief="flat", bd=0,
            activestyle="none",
            highlightthickness=0,
            exportselection=False)
        self._audio_lb.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)
        sb = tk.Scrollbar(self._list_container, command=self._audio_lb.yview,
                           orient="vertical")
        sb.grid(row=0, column=1, sticky="ns")
        self._audio_lb.config(yscrollcommand=sb.set)

        # Drop hint label
        self._hint_lbl = ctk.CTkLabel(
            self._left, text=f"⬇  {L['hint_drag']}",
            font=mk_font(12))
        self._hint_lbl.grid(row=3, column=0, padx=14, pady=(0, 8), sticky="w")

        # Count label
        self._lbl_count = ctk.CTkLabel(
            self._left, text="0 files",
            font=mk_font(11, bold=True))
        self._lbl_count.grid(row=4, column=0, padx=14, pady=(0, 10), sticky="w")

    # ── Right panel: settings + progress ──────────────────────────────────────
    def _build_right_panel(self):
        L = self._L()
        self._right = ctk.CTkFrame(self, corner_radius=0)
        self._right.grid(row=0, column=1, sticky="nsew", padx=(5, 10), pady=(70, 10))
        self._right.columnconfigure(0, weight=1)

        row = 0

        # Image section
        self._sec_img_lbl = ctk.CTkLabel(
            self._right, text=L["sec_image"],
            font=mk_font(14, bold=True), anchor="w")
        self._sec_img_lbl.grid(row=row, column=0, sticky="w", padx=14, pady=(14, 4))
        row += 1

        self._img_box = ctk.CTkFrame(self._right, height=72, corner_radius=10)
        self._img_box.grid(row=row, column=0, sticky="ew", padx=12, pady=(0, 8))
        self._img_box.grid_propagate(False)
        self._img_box.columnconfigure(0, weight=1)
        self._lbl_img = ctk.CTkLabel(
            self._img_box, text=L["no_img"],
            font=mk_font(12), anchor="w")
        self._lbl_img.grid(row=0, column=0, padx=12, pady=4, sticky="w")
        self._btn_img = ctk.CTkButton(
            self._img_box, text=L["btn_img"],
            font=mk_font(12, bold=True), height=30, corner_radius=8,
            command=self._select_image)
        self._btn_img.grid(row=1, column=0, padx=12, pady=(0, 8), sticky="w")
        row += 1

        self._div_r(row); row += 1

        # Output section
        self._sec_out_lbl = ctk.CTkLabel(
            self._right, text=L["sec_output"],
            font=mk_font(14, bold=True), anchor="w")
        self._sec_out_lbl.grid(row=row, column=0, sticky="w", padx=14, pady=(4, 4))
        row += 1
        self._lbl_out = ctk.CTkLabel(
            self._right, text=L["no_out"],
            font=mk_font(11), anchor="w", wraplength=280)
        self._lbl_out.grid(row=row, column=0, sticky="w", padx=14, pady=(0, 4))
        row += 1
        self._btn_out = ctk.CTkButton(
            self._right, text=L["btn_out"],
            font=mk_font(12, bold=True), height=32, corner_radius=8,
            command=self._select_output)
        self._btn_out.grid(row=row, column=0, padx=12, pady=(0, 8), sticky="w")
        row += 1

        self._div_r(row); row += 1

        # Settings
        self._sec_set_lbl = ctk.CTkLabel(
            self._right, text=L["sec_settings"],
            font=mk_font(14, bold=True), anchor="w")
        self._sec_set_lbl.grid(row=row, column=0, sticky="w", padx=14, pady=(4, 6))
        row += 1

        self._lbl_fname = ctk.CTkLabel(
            self._right, text=L["lbl_fname"],
            font=mk_font(12), anchor="w")
        self._lbl_fname.grid(row=row, column=0, sticky="w", padx=14, pady=(0, 2))
        row += 1
        self._entry_fname = ctk.CTkEntry(
            self._right, placeholder_text="Audio_Output",
            font=mk_font(13), height=36, corner_radius=8)
        self._entry_fname.insert(0, "Audio_Output")
        self._entry_fname.grid(row=row, column=0, padx=12, pady=(0, 8), sticky="ew")
        row += 1

        self._lbl_chunk = ctk.CTkLabel(
            self._right, text=L["lbl_chunk"],
            font=mk_font(12), anchor="w")
        self._lbl_chunk.grid(row=row, column=0, sticky="w", padx=14, pady=(0, 2))
        row += 1
        chunk_row = ctk.CTkFrame(self._right, fg_color="transparent")
        chunk_row.grid(row=row, column=0, padx=12, pady=(0, 8), sticky="ew")
        self._entry_chunk = ctk.CTkEntry(
            chunk_row, font=mk_font(14, bold=True),
            height=36, width=80, corner_radius=8, justify="center")
        self._entry_chunk.insert(0, "1")
        self._entry_chunk.pack(side="left")
        self._lbl_all = ctk.CTkLabel(
            chunk_row, text="  (หรือกรอก all = รวมทั้งหมด)",
            font=mk_font(11))
        self._lbl_all.pack(side="left")
        row += 1

        self._div_r(row); row += 1

        # Progress
        self._sec_prog_lbl = ctk.CTkLabel(
            self._right, text=L["sec_progress"],
            font=mk_font(14, bold=True), anchor="w")
        self._sec_prog_lbl.grid(row=row, column=0, sticky="w", padx=14, pady=(4, 6))
        row += 1

        self._prog_bar = ctk.CTkProgressBar(
            self._right, height=12, corner_radius=6)
        self._prog_bar.set(0)
        self._prog_bar.grid(row=row, column=0, padx=12, pady=(0, 6), sticky="ew")
        row += 1

        self._lbl_status = ctk.CTkLabel(
            self._right, text=L["status_idle"],
            font=mk_font(12, mono=True), anchor="w")
        self._lbl_status.grid(row=row, column=0, padx=14, pady=(0, 10), sticky="w")
        row += 1

        # Start button
        self._btn_start = ctk.CTkButton(
            self._right, text=L["btn_start"],
            font=mk_font(20, bold=True), height=56, corner_radius=14,
            command=self._start)
        self._btn_start.grid(row=row, column=0, padx=12, pady=(0, 16), sticky="ew")

    def _div_r(self, row):
        ctk.CTkFrame(self._right, height=1, corner_radius=0
                     ).grid(row=row, column=0, sticky="ew", padx=12, pady=6)

    # ── DnD ────────────────────────────────────────────────────────────────────
    def _init_dnd(self):
        if not DND_OK:
            return
        try:
            # Register Toplevel itself
            self.drop_target_register(DND_FILES)
            self.dnd_bind('<<Drop>>', self._on_drop)
            self.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self.dnd_bind('<<DragLeave>>', self._on_drag_leave)

            # Register raw tk Listbox widget
            lb = self._audio_lb
            lb.drop_target_register(DND_FILES)
            lb.dnd_bind('<<Drop>>', self._on_drop)

            # Register list_container frame
            self._list_container.drop_target_register(DND_FILES)
            self._list_container.dnd_bind('<<Drop>>', self._on_drop)

        except AttributeError:
            # If dnd_bind / drop_target_register not available on widgets,
            # try raw Tcl API (works when TkinterDnD._require was called on root)
            try:
                for w in [self, self._audio_lb, self._list_container]:
                    self.tk.call('tkdnd::drop_target', 'register', w, 'DND_Files')
                    w.bind('<<Drop>>', self._on_drop)
            except Exception as e2:
                print(f"Audio DnD fallback error: {e2}")
        except Exception as e:
            print(f"Audio DnD init error: {e}")

    def _on_drag_enter(self, e=None):
        T = self._T
        self._list_container.configure(highlightbackground=T["accent"],
                                        highlightthickness=2)

    def _on_drag_leave(self, e=None):
        T = self._T
        self._list_container.configure(highlightbackground=T["bg_border"],
                                        highlightthickness=1)

    def _on_drop(self, event):
        self._on_drag_leave()
        try:
            raw = event.data if hasattr(event, 'data') else ""
            files = self.tk.splitlist(raw)
        except Exception:
            return
        files = [f.strip("{}") for f in files]
        audios = [f for f in files if os.path.splitext(f)[1].lower() in ('.m4a', '.mp3', '.wav')]
        images = [f for f in files if os.path.splitext(f)[1].lower() in ('.jpg', '.jpeg', '.png')]
        if images:
            self.image_path = images[-1]
            self._lbl_img.configure(text=f"✅  {os.path.basename(self.image_path)}")
        if audios:
            existing = set(self.audio_paths)
            for f in audios:
                if f not in existing:
                    self.audio_paths.append(f)
                    existing.add(f)
            self._sort_natural()   # auto natural-sort on drop

    # ── Sort ────────────────────────────────────────────────────────────────────
    def _sort_natural(self):
        """Sort by numeric sequence in filename: ep1 < ep2 < ep10"""
        self.audio_paths.sort(key=_natural_key)
        self._refresh_list()

    def _sort_alpha(self):
        """Sort alphabetically A→Z"""
        self.audio_paths.sort(key=lambda p: os.path.basename(p).lower())
        self._refresh_list()

    def _refresh_list(self):
        self._audio_lb.delete(0, tk.END)
        T = self._T
        self._audio_lb.configure(bg=T["bg_input"], fg=T["fg_body"],
                                   selectbackground=T["accent"],
                                   selectforeground=T["fg_primary"])
        for i, fp in enumerate(self.audio_paths, 1):
            self._audio_lb.insert(tk.END, f"  {i:04d}  ·  {os.path.basename(fp)}")
        cnt = len(self.audio_paths)
        self._lbl_count.configure(
            text=self._L()["cnt_files"].format(n=cnt),
            text_color=self._T["success"] if cnt else self._T["fg_muted"])
        # auto-fill chunk entry
        if cnt:
            self._entry_chunk.delete(0, "end")
            self._entry_chunk.insert(0, str(cnt))

    # ── Controls ────────────────────────────────────────────────────────────────
    def _select_audio(self):
        fps = filedialog.askopenfilenames(
            title="เลือกไฟล์เสียง (Ctrl+Click เลือกหลายไฟล์)",
            filetypes=[("Audio", "*.m4a *.mp3 *.wav"), ("All", "*.*")])
        if fps:
            existing = set(self.audio_paths)
            for f in fps:
                if f not in existing:
                    self.audio_paths.append(f)
                    existing.add(f)
            self._sort_natural()

    def _delete_selected(self):
        idxs = list(self._audio_lb.curselection())[::-1]
        for i in idxs:
            self.audio_paths.pop(i)
        self._refresh_list()

    def _clear_audio(self):
        self.audio_paths.clear()
        self._refresh_list()

    def _select_image(self):
        f = filedialog.askopenfilename(
            title="เลือกรูปภาพพื้นหลัง",
            filetypes=[("Image", "*.jpg *.jpeg *.png"), ("All", "*.*")])
        if f:
            self.image_path = f
            self._lbl_img.configure(text=f"✅  {os.path.basename(f)}")

    def _select_output(self):
        d = filedialog.askdirectory(title="เลือกโฟลเดอร์บันทึกวิดีโอ")
        if d:
            self.output_dir = d
            self._lbl_out.configure(
                text=f"📂  {d}",
                text_color=self._T["warn"])

    def _toggle_lang(self):
        new = "EN" if self._lang == "TH" else "TH"
        set_lang(new)

    # ── Process ────────────────────────────────────────────────────────────────
    def _start(self):
        L = self._L()
        if not self.audio_paths or not self.image_path or not self.output_dir:
            messagebox.showwarning("Audio By Keawgood", L["err_missing"])
            return
        raw = self._entry_chunk.get().strip()
        if raw.lower() == "all":
            chunk = len(self.audio_paths)
        else:
            try:
                chunk = int(raw)
                if chunk <= 0:
                    raise ValueError
            except ValueError:
                messagebox.showwarning("Audio By Keawgood", L["err_num"])
                return
        if self._processing:
            return
        self._processing = True
        self._btn_start.configure(state="disabled")
        self._entry_chunk.configure(state="disabled")
        self._entry_fname.configure(state="disabled")
        self._prog_bar.set(0)
        threading.Thread(
            target=self._process,
            args=(chunk,),
            daemon=True).start()

    def _process(self, chunk_size: int):
        L = self._L()
        try:
            import moviepy.editor as mpe
        except ImportError:
            self.after(0, lambda: messagebox.showerror("Error", L["err_moviepy"]))
            self.after(0, self._done)
            return

        files   = list(self.audio_paths)  # snapshot
        base    = self._entry_fname.get().strip() or "Audio_Output"
        n_total = (len(files) + chunk_size - 1) // chunk_size

        for idx, start in enumerate(range(0, len(files), chunk_size), 1):
            chunk = files[start: start + chunk_size]
            self.after(0, lambda i=idx: (
                self._lbl_status.configure(
                    text=L["status_proc"].format(c=i, t=n_total)),
                self._prog_bar.set(i / n_total)
            ))
            try:
                clips  = [mpe.AudioFileClip(f) for f in chunk]
                audio  = mpe.concatenate_audioclips(clips)
                img    = mpe.ImageClip(self.image_path).set_duration(audio.duration).set_audio(audio)
                fname  = f"{base}.mp4" if n_total == 1 else f"{base}_Part{idx:03d}.mp4"
                out    = os.path.join(self.output_dir, fname)
                img.write_videofile(out, fps=1, codec="libx264", audio_codec="aac", logger=None)
                for c in clips: c.close()
                audio.close(); img.close()
            except Exception as e:
                print(f"Chunk {idx} error: {e}")

        self.after(0, self._done)

    def _done(self):
        L = self._L()
        self._lbl_status.configure(
            text=L["status_done"],
            text_color=self._T["success"])
        self._prog_bar.set(1.0)
        self._processing = False
        self._btn_start.configure(state="normal")
        self._entry_chunk.configure(state="normal")
        self._entry_fname.configure(state="normal")
        messagebox.showinfo("Audio By Keawgood", L["status_done"])

    # ── Refresh lang labels ─────────────────────────────────────────────────────
    def _refresh_lang(self):
        L = self._L()
        try:
            self.title(L["title"])
            self._lbl_title.configure(text=L["title"])
            self._lbl_sub.configure(text=L["sub"])
            self._btn_lang.configure(text=L["lang_switch"])
            self._sec_audio.configure(text=L["sec_audio"])
            self._sec_img_lbl.configure(text=L["sec_image"])
            self._sec_out_lbl.configure(text=L["sec_output"])
            self._sec_set_lbl.configure(text=L["sec_settings"])
            self._sec_prog_lbl.configure(text=L["sec_progress"])
            self._btn_add.configure(text=L["btn_add"])
            self._btn_clear.configure(text=L["btn_clear"])
            self._btn_sort_az.configure(text=L["btn_sort_az"])
            self._btn_sort_nat.configure(text=L["btn_sort_nat"])
            self._btn_img.configure(text=L["btn_img"])
            self._btn_out.configure(text=L["btn_out"])
            self._btn_start.configure(text=L["btn_start"])
            self._lbl_fname.configure(text=L["lbl_fname"])
            self._lbl_chunk.configure(text=L["lbl_chunk"])
            self._hint_lbl.configure(text=f"⬇  {L['hint_drag']}")
            self._lbl_status.configure(text=L["status_idle"])
        except Exception:
            pass

    # ── Theme apply ─────────────────────────────────────────────────────────────
    def _apply_theme(self):
        T = self._T
        self.configure(fg_color=T["bg_root"])
        for frame in [self._left, self._right, self._topbar]:
            try: frame.configure(fg_color=T["bg_card"])
            except Exception: pass

        self._lbl_title.configure(text_color=T["accent"])
        self._lbl_sub.configure(text_color=T["fg_muted"])

        self._list_container.configure(
            bg=T["bg_input"],
            highlightbackground=T["bg_border"])
        self._audio_lb.configure(
            bg=T["bg_input"], fg=T["fg_body"],
            selectbackground=T["accent"],
            selectforeground=T["fg_primary"])

        self._img_box.configure(fg_color=T["bg_input"])
        self._lbl_img.configure(text_color=T["fg_muted"])
        self._lbl_out.configure(text_color=T["fg_muted"])

        accent_btns = [self._btn_add, self._btn_sort_nat, self._btn_sort_az,
                       self._btn_img, self._btn_out, self._btn_start,
                       self._btn_lang]
        for b in accent_btns:
            b.configure(fg_color=T["accent"], hover_color=T["accent2"],
                        text_color=T["fg_primary"])

        self._btn_del_sel.configure(
            fg_color=T["error"], hover_color="#B91C1C",
            text_color="#FFFFFF")
        self._btn_clear.configure(
            fg_color=T["error"], hover_color="#B91C1C",
            text_color="#FFFFFF")

        self._prog_bar.configure(
            fg_color=T["bg_input"], progress_color=T["accent"])

        for lbl in [self._sec_audio, self._sec_img_lbl, self._sec_out_lbl,
                    self._sec_set_lbl, self._sec_prog_lbl]:
            lbl.configure(text_color=T["fg_primary"])

        for lbl in [self._lbl_count, self._hint_lbl, self._lbl_status,
                    self._lbl_fname, self._lbl_chunk, self._lbl_all]:
            lbl.configure(text_color=T["fg_muted"])

        self._entry_fname.configure(
            fg_color=T["bg_input"], border_color=T["bg_border"],
            text_color=T["fg_body"])
        self._entry_chunk.configure(
            fg_color=T["bg_input"], border_color=T["bg_border"],
            text_color=T["fg_body"])

# ══════════════════════════════════════════════════════════════════════════════

try:
    from curl_cffi import requests as cffi_requests
    HAS_CURL = True
except ImportError:
    import requests as cffi_requests  # type: ignore
    HAS_CURL = False

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

NV_LANG: dict = {
    "th": {
        "app_title": "📖 Novel By Keawgood", "subtitle": "โหลดนิยายอัตโนมัติ (Hybrid Mode)",
        "url_label": "URL หน้าแนะนำนิยาย หรือ สารบัญ", "url_placeholder": "ตัวอย่าง: https://twkan.com/book/61470.html",
        "fetch_btn": "🔍 ดึงรายการตอน", "fetching": "⏳ กำลังดึงข้อมูล…",
        "chapters_found": "พบทั้งหมด: {n} ตอน", "no_chapters": "— ยังไม่พบตอน —",
        "range_label": "ช่วงตอนที่ต้องการโหลด", "from_label": "จากตอนที่:", "to_label": "ถึงตอนที่:",
        "workers_label": "Workers (เธรด):", "delay_min": "Delay ต่ำสุด (วิ):", "delay_max": "Delay สูงสุด (วิ):",
        "headless_label": "ซ่อนหน้าต่าง Chrome", "naming_label": "รูปแบบชื่อไฟล์",
        "naming_hint": "ใช้ [n] แทนเลขตอน — เช่น  ตอนที่ [n]  หรือ  Chapter [n]",
        "save_label": "โฟลเดอร์บันทึก:", "browse_btn": "📁 เลือก",
        "start_btn": "⚡ เริ่มโหลด", "stop_btn": "⛔ หยุด", "clear_btn": "🗑 ล้าง Log",
        "log_label": "บันทึกกิจกรรม", "bypass_btn": "🛡 ขอ Cloudflare Cookie",
        "tip_workers": "💡 หมายเหตุ: ระบบป้องกันการข้ามตอน หากโหลดไม่ผ่านจะพยายามโหลดตอนเดิมซ้ำจนสำเร็จ",
        "lang_btn": "🌐 English", "theme_btn": "☀️ โหมดสว่าง",
        "err_no_url": "กรุณากรอก URL ก่อน", "err_no_fetch": "กรุณากด 'ดึงรายการตอน' ก่อนเริ่มโหลด",
        "err_no_dir": "กรุณาเลือกโฟลเดอร์บันทึก",
        "err_range": "ช่วงตอนไม่ถูกต้อง (ต้องอยู่ระหว่าง 1–{n})", "err_workers": "Workers ต้องเป็น 1–20",
        "err_delay": "Delay ต้องเป็นตัวเลขบวก และ min ≤ max",
        "done_msg": "✨ โหลดสำเร็จครบถ้วน! สำเร็จ {ok} | {t:.1f}s",
        "stopped_msg": "⛔ หยุดโดยผู้ใช้ — สำเร็จ {ok} | {t:.1f}s",
        "saved_at": "📂 ไฟล์บันทึกที่: {d}", "paste": "📋 วาง", "copy": "📄 คัดลอก",
    },
    "en": {
        "app_title": "📖 Novel By Keawgood", "subtitle": "Automatic Novel Downloader (Hybrid Mode)",
        "url_label": "Novel Info URL or TOC", "url_placeholder": "Example: https://twkan.com/book/61470.html",
        "fetch_btn": "🔍 Fetch Chapters", "fetching": "⏳ Fetching…",
        "chapters_found": "Found: {n} chapters", "no_chapters": "— No chapters found —",
        "range_label": "Chapter Range to Download", "from_label": "From chapter:", "to_label": "To chapter:",
        "workers_label": "Workers (threads):", "delay_min": "Delay min (s):", "delay_max": "Delay max (s):",
        "headless_label": "Hide Chrome window", "naming_label": "File Naming Pattern",
        "naming_hint": "Use [n] for chapter number — e.g.  Chapter [n]  or  ตอนที่ [n]",
        "save_label": "Save folder:", "browse_btn": "📁 Browse",
        "start_btn": "⚡ Start Download", "stop_btn": "⛔ Stop", "clear_btn": "🗑 Clear Log",
        "log_label": "Activity Log", "bypass_btn": "🛡 Get Cloudflare Cookie",
        "tip_workers": "💡 Note: Failed chapters will be retried infinitely until successful.",
        "lang_btn": "🌐 ภาษาไทย", "theme_btn": "☀️ Light Mode",
        "err_no_url": "Please enter a URL first", "err_no_fetch": "Please fetch chapters before starting download",
        "err_no_dir": "Please select a save folder",
        "err_range": "Invalid chapter range (must be 1–{n})", "err_workers": "Workers must be between 1–20",
        "err_delay": "Delay must be positive numbers and min ≤ max",
        "done_msg": "✨ Done! Success {ok} | {t:.1f}s", "stopped_msg": "⛔ Stopped by user — Success {ok} | {t:.1f}s",
        "saved_at": "📂 Files saved at: {d}", "paste": "📋 Paste", "copy": "📄 Copy",
    },
}

NV_THEMES: dict = {
    "dark": {
        "BG": "#0f111a", "PANEL": "#1a1d2d", "ACCENT": "#6a5acd",
        "ACCENT_H": "#7b68ee", "ACCENT2": "#20c997", "FG": "#e2e8f0",
        "FG_DIM": "#94a3b8", "SUCCESS": "#20c997", "ERROR": "#f43f5e",
        "WARN": "#f59e0b", "ENTRY_BG": "#0b0c13", "BORDER": "#2e344e",
    },
    "light": {
        "BG": "#f8fafc", "PANEL": "#ffffff", "ACCENT": "#4f46e5",
        "ACCENT_H": "#6366f1", "ACCENT2": "#059669", "FG": "#0f172a",
        "FG_DIM": "#64748b", "SUCCESS": "#059669", "ERROR": "#e11d48",
        "WARN": "#d97706", "ENTRY_BG": "#f1f5f9", "BORDER": "#e2e8f0",
    },
}

_nv_cf_lock = threading.Lock()
_nv_cf_cookies: list = []
_nv_cf_user_agent: str = ""
NV_USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
]

NV_SITE_RULES: dict = {
    "twkan.com": {
        "link_filter": lambda abs_url, bid: re.search(r'\d+\.html$', abs_url) and "index.html" not in abs_url,
        "content_selector": ("div", {"id": re.compile(r'(content|chaptercontent|txtcontent0|BookText|read)', re.I)}),
        "unwanted_tags": ["script", "style", "ins", "h1"],
        "encoding": "utf-8",
    },
    "default": {
        "link_filter": lambda abs_url, bid: len(abs_url) > 10 and not abs_url.endswith("#") and (
            re.search(r'(chapter|chap|vol|part|/p/|/read/|/txt/|\d{3,}\.html|\d+/?$)', abs_url.lower())
            or (bid and bid in abs_url)),
        "content_selector": ("div", {"class": re.compile(r'(content|chapter|text|read|body|main|entry|post|article)', re.I)}),
        "unwanted_tags": ["script", "style", "ins", "aside", "nav", "header", "footer", "iframe"],
        "encoding": None,
    },
}

_NV_STEALTH_JS = """
Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
Object.defineProperty(navigator, 'languages', { get: () => ['en-US', 'en'] });
Object.defineProperty(navigator, 'platform', { get: () => 'Win32' });
"""

_NV_AD_PATTERNS = re.compile(
    r'(advertisement|sponsored|subscribe|follow us|please support|read at|'
    r'translator.?note|t/n:|tl.?note|visit .{0,30} for .{0,30} chapters|'
    r'patreon\.com|discord\.gg|ko-?fi\.com|www\.\w+\.com|'
    r'手机版阅读|最新网址|更新最快|下载APP|无弹窗|顶点小说|QQ群|微信号|公众号|一秒记住|'
    r'推荐本书|txt下载|求推荐|求收藏|章节错误|点此举报|加入书签|上一章|返回目录|下一章|'
    r'天才一秒记住本站地址|twkan|台湾看小说|台灣小說網|₮₩₭₳₦|請記住|觀看最快|章節更新|loadAdv)',
    re.I
)


def nv_playwright_get_cookies(url: str, log_fn, headless: bool = True) -> bool:
    global _nv_cf_cookies, _nv_cf_user_agent
    if not HAS_PLAYWRIGHT:
        log_fn("⚠️ Playwright not installed — skipping browser bypass"); return False
    log_fn("🌐 Opening stealth browser to obtain Cloudflare clearance…")
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=headless, args=["--disable-blink-features=AutomationControlled"])
            ctx = browser.new_context(user_agent=random.choice(NV_USER_AGENTS), viewport={"width": 1366, "height": 768})
            ctx.add_init_script(_NV_STEALTH_JS)
            page = ctx.new_page()
            try:
                page.goto(url, wait_until="domcontentloaded", timeout=30_000)
            except Exception:
                pass
            for _ in range(20):
                if "just a moment" not in page.title().lower() and "cloudflare" not in page.title().lower():
                    break
                time.sleep(1)
            _nv_cf_cookies = ctx.cookies()
            _nv_cf_user_agent = page.evaluate("navigator.userAgent")
            browser.close()
        log_fn("✅ Cookie obtained!" if _nv_cf_cookies else "❌ No cookies obtained")
        return bool(_nv_cf_cookies)
    except Exception as e:
        log_fn(f"❌ Playwright error: {e}"); return False


def nv_make_session():
    ua = _nv_cf_user_agent or random.choice(NV_USER_AGENTS)
    sess = cffi_requests.Session(impersonate="chrome120") if HAS_CURL else cffi_requests.Session()
    sess.headers.update({"User-Agent": ua, "Accept-Language": "en-US,en;q=0.9"})
    for c in _nv_cf_cookies:
        try:
            sess.cookies.set(c["name"], c["value"], domain=c.get("domain", ""))
        except Exception:
            pass
    return sess


def nv_decode_response(content: bytes, hint=None) -> str:
    if hint:
        try: return content.decode(hint, errors="replace")
        except Exception: pass
    if HAS_CHARDET:
        enc = chardet.detect(content[:4096]).get("encoding") or "utf-8"
        try: return content.decode(enc, errors="replace")
        except Exception: pass
    for enc in ("utf-8", "gbk", "big5"):
        try: return content.decode(enc, errors="replace")
        except Exception: continue
    return content.decode("utf-8", errors="replace")


def nv_detect_site(url: str) -> dict:
    for domain, rules in NV_SITE_RULES.items():
        if domain != "default" and domain in url:
            return rules
    return NV_SITE_RULES["default"]


def nv_extract_book_id(url: str) -> str:
    for pat in [r'/book/(\w+)', r'/b/(\w+)', r'/n/(\w+)', r'/(\d{4,})/?(?:\?|$|#)']:
        m = re.search(pat, url)
        if m: return m.group(1)
    parts = url.rstrip("/").split("/")
    return re.sub(r'[^\w]', '', parts[-1] if parts else "") or ""


def nv_fetch_toc(catalog_url: str, log_fn=print) -> list:
    if not HAS_BS4:
        log_fn("❌ beautifulsoup4 ไม่ได้ติดตั้ง — pip install beautifulsoup4"); return []
    rules = nv_detect_site(catalog_url)
    book_id = nv_extract_book_id(catalog_url)
    session = nv_make_session()
    html_text = None
    for _ in range(3):
        try:
            resp = session.get(catalog_url, timeout=20)
            if resp.status_code == 200:
                html_text = nv_decode_response(resp.content, rules.get("encoding")); break
        except Exception:
            time.sleep(2)
    if html_text is None:
        with _nv_cf_lock:
            if not _nv_cf_cookies:
                nv_playwright_get_cookies(catalog_url, log_fn)
        session = nv_make_session()
        for attempt in range(4):
            try:
                resp = session.get(catalog_url, timeout=20)
                if resp.status_code == 200:
                    html_text = nv_decode_response(resp.content, rules.get("encoding")); break
                time.sleep(2 ** attempt)
            except Exception:
                time.sleep(2)
        else:
            return []
    if not html_text: return []
    soup = BeautifulSoup(html_text, "html.parser")
    links = []
    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        abs_url = urljoin(catalog_url, href)
        try:
            if rules["link_filter"](abs_url, book_id):
                links.append(abs_url)
        except Exception:
            pass
    unique = list(dict.fromkeys(links))
    try:
        unique.sort(key=lambda u: int(re.search(r'/(\d+)(?:\.html?)?/?(?:\?.*)?$', u).group(1)))
    except Exception:
        pass
    if not unique:
        log_fn("⚠️ หาลิงก์สารบัญจากหน้านี้ไม่พบ (อาจต้องใช้โหมด 'ไต่ลิงก์ทีละตอน')")
        return [catalog_url]
    return unique


def nv_clean_text(div, unwanted_tags: list) -> str:
    for bad in div.find_all(unwanted_tags):
        bad.decompose()
    html_content = re.sub(r'<br\s*/?>', '\n', str(div))
    clean_soup = BeautifulSoup(html_content, "html.parser")
    lines = []
    for ln in clean_soup.get_text(separator="\n").splitlines():
        s = ln.strip()
        if s and not _NV_AD_PATTERNS.search(s):
            lines.append(s)
    return "\n".join(lines).strip()


def nv_fetch_chapter(session, chapter_num, url, rules, save_dir, stop_event, naming_pattern, log_fn, dmin, dmax) -> str:
    attempt = 0
    while not stop_event.is_set():
        if attempt > 0:
            stop_event.wait(min(30, (2 ** min(attempt, 5))) + random.uniform(1.0, 3.0))
        else:
            stop_event.wait(random.uniform(dmin, dmax))
        if stop_event.is_set(): return f"⛔ [{chapter_num:04d}] ยกเลิก"
        attempt += 1
        try:
            resp = session.get(url, timeout=30, headers={"Referer": url})
            if resp.status_code == 200:
                soup = BeautifulSoup(nv_decode_response(resp.content, rules.get("encoding")), "html.parser")
                title = soup.find("h1")
                title_text = title.get_text(strip=True) if title else f"Chapter {chapter_num:04d}"
                tag, attrs = rules["content_selector"]
                div = soup.find(tag, attrs)
                if not div:
                    for pid in ["content", "txtcontent0", "chaptercontent", "BookText", "chapter-content", "txtContent"]:
                        div = soup.find("div", id=re.compile(pid, re.I))
                        if div: break
                if not div:
                    all_divs = soup.find_all("div")
                    if all_divs:
                        best_div = max(all_divs, key=lambda d: len(d.get_text(strip=True)))
                        if len(best_div.get_text(strip=True)) > 200: div = best_div
                if not div:
                    log_fn(f"⚠️ [{chapter_num:04d}] ไม่พบเนื้อหา จะพยายามโหลดใหม่..."); continue
                text = nv_clean_text(div, rules.get("unwanted_tags", []))
                if len(text) < 30:
                    log_fn(f"⚠️ [{chapter_num:04d}] เนื้อหาสั้นผิดปกติ จะพยายามโหลดใหม่..."); continue
                safe_title = re.sub(r'[\\/*?:"<>|\r\n]', "", title_text)[:80]
                name = (re.sub(r'[\\/*?:"<>|]', "", naming_pattern.replace("[n]", str(chapter_num)))
                        if naming_pattern.strip() else f"{chapter_num:04d}_{safe_title}")
                with open(os.path.join(save_dir, name.strip() + ".txt"), "w", encoding="utf-8") as f:
                    f.write(f"{title_text}\n\n{text}\n\n(本集结束)")
                return f"✅ [{chapter_num:04d}] {title_text[:60]}"
            elif resp.status_code in (403, 429, 503):
                log_fn(f"⚠️ [{chapter_num:04d}] ติด Block ({resp.status_code}) จะพยายามโหลดใหม่...")
            else:
                log_fn(f"⚠️ [{chapter_num:04d}] HTTP {resp.status_code} จะพยายามโหลดใหม่...")
        except Exception:
            log_fn(f"⚠️ [{chapter_num:04d}] Error: เชื่อมต่อขัดข้อง จะพยายามโหลดใหม่...")
    return f"⛔ [{chapter_num:04d}] ยกเลิก"


class NovelByKeawgoodWindow(tk.Toplevel):
    """Novel V2 — runs as a Toplevel inside the CTk hub."""

    def __init__(self, master):
        super().__init__(master)
        self._lang       = "th"
        self._theme_name = "dark"
        self._theme      = NV_THEMES["dark"]
        self._links: list = []
        self._running    = False
        self._stop_event = threading.Event()

        self.title(NV_LANG["th"]["app_title"])
        self.geometry("950x880")
        self.minsize(800, 680)
        self.resizable(True, True)

        self.font_family = "TH Sarabun PSK"
        self.font_h1   = tkfont.Font(family=self.font_family, size=28, weight="bold")
        self.font_h2   = tkfont.Font(family=self.font_family, size=20, weight="bold")
        self.font_body = tkfont.Font(family=self.font_family, size=18)
        self.font_btn  = tkfont.Font(family=self.font_family, size=18, weight="bold")
        self.font_log  = tkfont.Font(family=self.font_family, size=16)

        self._apply_theme()
        self._setup_global_clipboard_bindings()
        self._build()
        self.lift()
        self.focus_force()

    def _adjust_font_size(self, delta):
        for f in [self.font_h1, self.font_h2, self.font_body, self.font_btn, self.font_log]:
            f.configure(size=max(10, f.cget("size") + delta))

    def _setup_global_clipboard_bindings(self):
        self.bind_all("<KeyPress>", self._handle_global_shortcuts)

    def _handle_global_shortcuts(self, event):
        if not (event.state & 0x0004 or event.state & 0x0008 or event.state & 0x20000): return
        char = getattr(event, "char", "").lower()
        keysym = getattr(event, "keysym", "").lower()
        if char in ("c", "แ") or keysym in ("c", "oae"): return self._execute_copy(event)
        elif char in ("v", "อ") or keysym in ("v", "oang"): return self._execute_paste(event)
        elif char in ("x", "ป") or keysym in ("x", "porpla"): return self._execute_cut(event)
        elif char in ("a", "ฟ") or keysym in ("a", "forfan"): return self._execute_select_all(event)

    def _execute_copy(self, event):
        w = event.widget
        try:
            text = w.selection_get() if isinstance(w, (tk.Entry, ttk.Entry)) else w.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.clipboard_clear(); self.clipboard_append(text); return "break"
        except tk.TclError: pass

    def _execute_paste(self, event):
        w = event.widget
        try:
            text = self.clipboard_get()
            if isinstance(w, (tk.Entry, ttk.Entry)):
                try: w.delete(tk.SEL_FIRST, tk.SEL_LAST)
                except tk.TclError: pass
                w.insert(tk.INSERT, text); return "break"
            elif isinstance(w, tk.Text) and w.cget("state") == "normal":
                w.insert(tk.INSERT, text); return "break"
        except tk.TclError: pass

    def _execute_cut(self, event):
        w = event.widget
        try:
            if isinstance(w, (tk.Entry, ttk.Entry)):
                text = w.selection_get()
                self.clipboard_clear(); self.clipboard_append(text)
                w.delete(tk.SEL_FIRST, tk.SEL_LAST); return "break"
        except tk.TclError: pass

    def _execute_select_all(self, event):
        w = event.widget
        if isinstance(w, (tk.Entry, ttk.Entry)):
            w.select_range(0, tk.END); w.icursor(tk.END); return "break"
        elif isinstance(w, tk.Text):
            w.tag_add(tk.SEL, "1.0", tk.END); return "break"

    def t(self, key: str, **kw) -> str:
        s = NV_LANG[self._lang].get(key, key)
        return s.format(**kw) if kw else s

    @property
    def T(self): return self._theme

    def _apply_theme(self):
        self._theme = NV_THEMES[self._theme_name]
        self.configure(bg=self.T["BG"])

    def _toggle_theme(self):
        self._theme_name = "light" if self._theme_name == "dark" else "dark"
        self._apply_theme()
        for w in self.winfo_children(): w.destroy()
        self._build()

    def _toggle_lang(self):
        self._lang = "en" if self._lang == "th" else "th"
        for w in self.winfo_children(): w.destroy()
        self._build()

    def _ttk_style(self):
        s = ttk.Style(self)
        s.theme_use("clam")
        s.configure("Accent.Horizontal.TProgressbar", troughcolor=self.T["ENTRY_BG"],
                    background=self.T["ACCENT2"], bordercolor=self.T["ENTRY_BG"], relief="flat")

    def _card(self, parent):
        return tk.Frame(parent, bg=self.T["PANEL"], bd=0, highlightthickness=1,
                        highlightbackground=self.T["BORDER"], padx=20, pady=16)

    def _btn(self, parent, key, cmd, bg=None, fg=None, hover_bg=None, width=None, text=None, custom_font=None):
        label = text if text is not None else self.t(key)
        btn = tk.Button(parent, text=label, command=cmd, bg=bg or self.T["ACCENT"], fg=fg or "white",
                        font=custom_font or self.font_btn, relief="flat", cursor="hand2",
                        activebackground=hover_bg or self.T["ACCENT_H"], activeforeground=fg or "white",
                        padx=12, pady=6, bd=0)
        if width: btn.config(width=width)
        return btn

    def _entry(self, parent, textvariable, width=None, justify="left", ipady=4):
        kw = dict(textvariable=textvariable, bg=self.T["ENTRY_BG"], fg=self.T["FG"],
                  insertbackground=self.T["FG"], relief="flat", font=self.font_body, bd=6, justify=justify)
        if width: kw["width"] = width
        return tk.Entry(parent, **kw)

    def _label(self, parent, key, font=None, fg=None, text=None):
        return tk.Label(parent, text=text if text is not None else self.t(key),
                        font=font or self.font_body, bg=self.T["PANEL"], fg=fg or self.T["FG"])

    def _build(self):
        self._ttk_style()
        hdr = tk.Frame(self, bg=self.T["BG"])
        hdr.pack(fill="x", padx=24, pady=(15, 5))
        left = tk.Frame(hdr, bg=self.T["BG"])
        left.pack(side="left", fill="y")
        tk.Label(left, text=self.t("app_title"), font=self.font_h1, bg=self.T["BG"], fg=self.T["ACCENT"]).pack(anchor="w")
        tk.Label(left, text=self.t("subtitle"), font=self.font_body, bg=self.T["BG"], fg=self.T["FG_DIM"]).pack(anchor="w")
        right = tk.Frame(hdr, bg=self.T["BG"])
        right.pack(side="right", fill="y")
        ctrl_row = tk.Frame(right, bg=self.T["BG"])
        ctrl_row.pack(anchor="e")
        tk.Button(ctrl_row, text="A-", command=lambda: self._adjust_font_size(-2), bg=self.T["BORDER"], fg=self.T["FG"],
                  font=("Arial", 10, "bold"), relief="flat", cursor="hand2", padx=8, pady=4, bd=0).pack(side="left", padx=(0, 4))
        tk.Button(ctrl_row, text="A+", command=lambda: self._adjust_font_size(2), bg=self.T["BORDER"], fg=self.T["FG"],
                  font=("Arial", 10, "bold"), relief="flat", cursor="hand2", padx=8, pady=4, bd=0).pack(side="left", padx=(0, 12))
        theme_icon = "🌙" if self._theme_name == "dark" else "☀️"
        tk.Button(ctrl_row, text=theme_icon, command=self._toggle_theme, bg=self.T["BORDER"], fg=self.T["FG"],
                  font=("Arial", 10), relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left", padx=(0, 8))
        tk.Button(ctrl_row, text=self.t("lang_btn"), command=self._toggle_lang, bg=self.T["BORDER"], fg=self.T["FG"],
                  font=self.font_btn, relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left")

        uc = self._card(self)
        uc.pack(fill="x", padx=24, pady=8)
        uc.columnconfigure(0, weight=1)
        self._label(uc, "url_label", font=self.font_h2).grid(row=0, column=0, sticky="w", pady=(0, 4))
        ur = tk.Frame(uc, bg=self.T["PANEL"])
        ur.grid(row=1, column=0, sticky="ew")
        ur.columnconfigure(0, weight=1)
        self.url_var = tk.StringVar(value="https://twkan.com/book/61470.html")
        self.url_entry = self._entry(ur, self.url_var, ipady=6)
        self.url_entry.grid(row=0, column=0, sticky="ew")
        btn_action = tk.Frame(ur, bg=self.T["PANEL"])
        btn_action.grid(row=0, column=1, padx=(10, 10))
        tk.Button(btn_action, text=self.t("paste"), command=self._ui_paste_url,
                  bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_body,
                  relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left", padx=2)
        tk.Button(btn_action, text=self.t("copy"), command=self._ui_copy_url,
                  bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_body,
                  relief="flat", cursor="hand2", padx=10, pady=4, bd=0).pack(side="left", padx=2)
        self.fetch_btn = self._btn(ur, "fetch_btn", self._on_fetch)
        self.fetch_btn.grid(row=0, column=2)
        self.count_lbl = tk.Label(uc, text=self.t("no_chapters"), font=self.font_btn,
                                  bg=self.T["PANEL"], fg=self.T["ACCENT2"])
        self.count_lbl.grid(row=2, column=0, sticky="w", pady=(8, 0))
        mr = tk.Frame(uc, bg=self.T["PANEL"])
        mr.grid(row=3, column=0, sticky="w", pady=(5, 0))
        self.scrape_mode = tk.StringVar(value="crawler")
        tk.Label(mr, text="ระบบดึงข้อมูล:", font=self.font_btn, bg=self.T["PANEL"], fg=self.T["FG"]).pack(side="left")
        tk.Radiobutton(mr, text="ดึงจากสารบัญ (โหลดพร้อมกันหลายไฟล์)", variable=self.scrape_mode, value="concurrent",
                       font=self.font_body, bg=self.T["PANEL"], fg=self.T["ACCENT2"],
                       selectcolor=self.T["ENTRY_BG"], activebackground=self.T["PANEL"]).pack(side="left", padx=10)
        tk.Radiobutton(mr, text="ไต่ลิงก์ทีละตอน (🌟ชัวร์ 100% สำหรับ TWKAN)", variable=self.scrape_mode, value="crawler",
                       font=self.font_body, bg=self.T["PANEL"], fg=self.T["WARN"],
                       selectcolor=self.T["ENTRY_BG"], activebackground=self.T["PANEL"]).pack(side="left", padx=10)

        cc = self._card(self)
        cc.pack(fill="x", padx=24, pady=8)
        for c in range(8): cc.columnconfigure(c, weight=1)
        rows_cfg = [
            ("from_label",    "start_var",     "1",   0, 0),
            ("to_label",      "end_var",        "10",  0, 2),
            ("workers_label", "workers_var",   "5",   0, 4),
            ("delay_min",     "delay_min_var", "1.5", 1, 0),
            ("delay_max",     "delay_max_var", "4.0", 1, 2),
        ]
        for lkey, attr, default, row, col in rows_cfg:
            self._label(cc, lkey).grid(row=row, column=col, sticky="w", pady=(0 if row == 0 else 12, 0))
            var = tk.StringVar(value=default)
            setattr(self, attr, var)
            self._entry(cc, var, width=6 if attr == "workers_var" else 8, justify="center"
                        ).grid(row=row, column=col+1, padx=(4, 16), ipady=2, pady=(0 if row == 0 else 12, 0))
        self.headless_var = tk.BooleanVar(value=True)
        tk.Checkbutton(cc, text=self.t("headless_label"), variable=self.headless_var,
                       font=self.font_body, bg=self.T["PANEL"], fg=self.T["FG_DIM"],
                       selectcolor=self.T["ENTRY_BG"], activebackground=self.T["PANEL"]
                       ).grid(row=0, column=6, columnspan=2, sticky="w", padx=(8, 0))
        self._btn(cc, "bypass_btn", self._on_bypass, bg=self.T["WARN"], fg="#1a1d2e", hover_bg="#fcd34d"
                  ).grid(row=1, column=4, columnspan=4, sticky="e", pady=(12, 0))
        tk.Frame(cc, bg=self.T["BORDER"], height=1).grid(row=2, column=0, columnspan=8, sticky="ew", pady=(16, 12))
        self._label(cc, "naming_label", font=self.font_h2).grid(row=3, column=0, columnspan=8, sticky="w")
        self.naming_var = tk.StringVar(value="ตอนที่ [n]" if self._lang == "th" else "Chapter [n]")
        nr = tk.Frame(cc, bg=self.T["PANEL"])
        nr.grid(row=4, column=0, columnspan=8, sticky="ew", pady=(8, 0))
        nr.columnconfigure(0, weight=1)
        self._entry(nr, self.naming_var, ipady=4).grid(row=0, column=0, sticky="ew")
        for i, preset in enumerate(["ตอนที่ [n]", "Chapter [n]", "第[n]章"]):
            tk.Button(nr, text=preset, command=lambda p=preset: self.naming_var.set(p),
                      bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_body,
                      relief="flat", cursor="hand2", padx=10, pady=4, bd=0).grid(row=0, column=i+1, padx=(8, 0))
        tk.Frame(cc, bg=self.T["BORDER"], height=1).grid(row=5, column=0, columnspan=8, sticky="ew", pady=(16, 12))
        dr = tk.Frame(cc, bg=self.T["PANEL"])
        dr.grid(row=6, column=0, columnspan=8, sticky="ew")
        dr.columnconfigure(1, weight=1)
        self._label(dr, "save_label").grid(row=0, column=0, sticky="w")
        self.save_dir_var = tk.StringVar(value=str(Path.home() / "Downloads" / "novels"))
        self._entry(dr, self.save_dir_var, ipady=4).grid(row=0, column=1, sticky="ew", padx=(10, 10))
        self._btn(dr, "browse_btn", self._browse, bg=self.T["ENTRY_BG"], fg=self.T["FG"]).grid(row=0, column=2)
        tk.Label(cc, text=self.t("tip_workers"), font=self.font_body, bg=self.T["PANEL"],
                 fg=self.T["WARN"]).grid(row=7, column=0, columnspan=8, sticky="w", pady=(10, 0))

        pg = tk.Frame(self, bg=self.T["BG"])
        pg.pack(fill="x", padx=24, pady=(8, 4))
        pg.columnconfigure(0, weight=1)
        self.prog_var = tk.DoubleVar(value=0)
        ttk.Progressbar(pg, variable=self.prog_var, maximum=100,
                        style="Accent.Horizontal.TProgressbar").grid(row=0, column=0, sticky="ew")
        self.prog_lbl = tk.Label(pg, text="0 / 0", font=self.font_btn, bg=self.T["BG"], fg=self.T["FG_DIM"])
        self.prog_lbl.grid(row=0, column=1, padx=(12, 0))

        br = tk.Frame(self, bg=self.T["BG"])
        br.pack(pady=8)
        self.start_btn = self._btn(br, "start_btn", self._on_start, bg=self.T["ACCENT2"], fg="#0f111a",
                                   hover_bg="#34d399", width=18)
        self.start_btn.pack(side="left", padx=8)
        self.stop_btn = self._btn(br, "stop_btn", self._on_stop, bg=self.T["ERROR"], fg="white",
                                  hover_bg="#fb7185", width=10)
        self.stop_btn.pack(side="left", padx=8)
        self.stop_btn.config(state="disabled")
        self._btn(br, "clear_btn", self._clear_log, bg=self.T["BORDER"], fg=self.T["FG"],
                  hover_bg=self.T["ENTRY_BG"], width=10).pack(side="left", padx=8)

        lf = self._card(self)
        lf.pack(fill="both", expand=True, padx=24, pady=(4, 20))
        lf.columnconfigure(0, weight=1)
        lf.rowconfigure(1, weight=1)
        self._label(lf, "log_label", font=self.font_h2).grid(row=0, column=0, sticky="w", pady=(0, 8))
        self.log = tk.Text(lf, bg=self.T["ENTRY_BG"], fg=self.T["FG"], font=self.font_log,
                           relief="flat", wrap="word", state="disabled", bd=10)
        self.log.grid(row=1, column=0, sticky="nsew")
        sb = tk.Scrollbar(lf, command=self.log.yview, bg=self.T["PANEL"], troughcolor=self.T["ENTRY_BG"])
        sb.grid(row=1, column=1, sticky="ns", padx=(4, 0))
        self.log.config(yscrollcommand=sb.set)
        for tag, col in [("ok", self.T["SUCCESS"]), ("err", self.T["ERROR"]),
                         ("warn", self.T["WARN"]), ("info", self.T["FG_DIM"]), ("head", self.T["ACCENT_H"])]:
            self.log.tag_config(tag, foreground=col)
        self._log("Novel By Keawgood  v5.0  (Hybrid TWKAN Edition) — ready ✓", "head")

    def _ui_paste_url(self):
        try:
            text = self.clipboard_get()
            self.url_entry.delete(0, tk.END)
            self.url_entry.insert(0, text)
        except tk.TclError: pass

    def _ui_copy_url(self):
        text = self.url_entry.get()
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)

    def _log(self, msg: str, tag: str = "info"):
        def _do():
            self.log.config(state="normal")
            self.log.insert("end", f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n", tag)
            self.log.see("end")
            self.log.config(state="disabled")
        self.after(0, _do)

    def _clear_log(self):
        self.log.config(state="normal")
        self.log.delete("1.0", "end")
        self.log.config(state="disabled")

    def _browse(self):
        if d := filedialog.askdirectory(title="Select Folder"):
            self.save_dir_var.set(d)

    def _on_bypass(self):
        if not (url := self.url_var.get().strip()):
            return messagebox.showerror("Error", self.t("err_no_url"))
        self._log("🛡 Requesting Cloudflare clearance manually…", "warn")
        threading.Thread(target=lambda: nv_playwright_get_cookies(
            url, lambda m: self._log(m), self.headless_var.get()), daemon=True).start()

    def _on_fetch(self):
        if not (url := self.url_var.get().strip()):
            return messagebox.showerror("Error", self.t("err_no_url"))
        if self.scrape_mode.get() == "crawler":
            self._log("💡 อยู่ใน 'โหมดไต่ลิงก์' ไม่จำเป็นต้องดึงสารบัญ สามารถกดปุ่ม [เริ่มโหลด] ได้เลย", "head")
            return
        self.fetch_btn.config(state="disabled", text=self.t("fetching"))
        self._log(f"🔍 {self.t('fetch_btn')}: {url}", "head")

        def _worker():
            self._links = nv_fetch_toc(url, log_fn=lambda m: self._log(m, "warn"))
            def _upd():
                n = len(self._links)
                self.count_lbl.config(text=self.t("chapters_found", n=n) if n else self.t("no_chapters"),
                                      fg=self.T["SUCCESS"] if n else self.T["ERROR"])
                if n: self.end_var.set(str(n))
                self.fetch_btn.config(state="normal", text=self.t("fetch_btn"))
            self.after(0, _upd)
        threading.Thread(target=_worker, daemon=True).start()

    def _on_start(self):
        if self._running: return
        if not (save_dir := self.save_dir_var.get().strip()):
            return messagebox.showerror("Error", self.t("err_no_dir"))
        os.makedirs(save_dir, exist_ok=True)
        mode = self.scrape_mode.get()
        if mode == "crawler":
            if not self.url_var.get().strip():
                return messagebox.showerror("Error", self.t("err_no_url"))
            self._running = True
            self._stop_event.clear()
            self.start_btn.config(state="disabled")
            self.stop_btn.config(state="normal")
            threading.Thread(target=self._crawler_worker, daemon=True).start()
            return
        if not self._links:
            return messagebox.showerror("Error", self.t("err_no_fetch"))
        try:
            s = int(self.start_var.get())
            e = int(self.end_var.get())
            workers = int(self.workers_var.get())
            dmin = float(self.delay_min_var.get())
            dmax = float(self.delay_max_var.get())
            assert 1 <= s <= e <= len(self._links) and 1 <= workers <= 20 and 0 <= dmin <= dmax
        except Exception:
            return messagebox.showerror("Input Error", "Check numbers (range, workers, delays)")
        self._running = True
        naming = self.naming_var.get().strip()
        self._stop_event.clear()
        self.start_btn.config(state="disabled")
        self.stop_btn.config(state="normal")
        selected = self._links[s - 1: e]
        self.prog_var.set(0)
        self.prog_lbl.config(text=f"0 / {len(selected)}")
        rules = nv_detect_site(self.url_var.get())
        self._log(f"⚡ Downloading {len(selected)} chapters (Concurrent Mode)...", "head")

        def _worker():
            session = nv_make_session()
            completed = 0
            t0 = time.time()
            log_wrapper = lambda m: self._log(m, "warn")
            with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as ex:
                futs = {ex.submit(nv_fetch_chapter, session, s + i, u, rules, save_dir,
                                  self._stop_event, naming, log_wrapper, dmin, dmax): s + i
                        for i, u in enumerate(selected) if not self._stop_event.is_set()}
                for fut in concurrent.futures.as_completed(futs):
                    if self._stop_event.is_set(): break
                    res = fut.result()
                    if "✅" in res: completed += 1; self._log(res, "ok")
                    elif "⛔" in res: self._log(res, "warn")
                    self.after(0, lambda c=completed: (
                        self.prog_var.set(c / len(selected) * 100),
                        self.prog_lbl.config(text=f"{c} / {len(selected)}")))
            self.after(0, lambda: self._on_done(completed, time.time() - t0))
        threading.Thread(target=_worker, daemon=True).start()

    def _crawler_worker(self):
        url = self.url_var.get().strip()
        save_dir = self.save_dir_var.get().strip()
        naming = self.naming_var.get().strip()
        rules = nv_detect_site(url)
        session = nv_make_session()
        current_url = url
        try: chapter_num = int(self.start_var.get())
        except Exception: chapter_num = 1
        completed = 0
        try: dmin, dmax = float(self.delay_min_var.get()), float(self.delay_max_var.get())
        except Exception: dmin, dmax = 1.5, 4.0
        self._log("═" * 52, "head")
        self._log(f"⚡ เริ่ม [โหมดไต่ลิงก์] จาก: {url}", "head")
        self.after(0, lambda: (self.prog_lbl.config(text="โหมดไต่ลิงก์: กำลังเริ่มต้น..."), self.prog_var.set(0)))
        t0 = time.time()
        while not self._stop_event.is_set():
            attempt, success = 0, False
            while attempt < 5 and not self._stop_event.is_set():
                if attempt > 0:
                    self._stop_event.wait(min(30, (2 ** attempt)) + random.uniform(1, 3))
                else:
                    self._stop_event.wait(random.uniform(dmin, dmax))
                attempt += 1
                try:
                    if not HAS_BS4: self._log("❌ beautifulsoup4 ไม่ได้ติดตั้ง", "err"); self._stop_event.set(); break
                    resp = session.get(current_url, timeout=30, headers={"Referer": current_url})
                    if resp.status_code == 200:
                        soup = BeautifulSoup(nv_decode_response(resp.content, rules.get("encoding")), "html.parser")
                        tag, attrs = rules.get("content_selector", ("div", {}))
                        div = soup.find(tag, attrs)
                        if not div:
                            for pid in ["txtcontent0", "content", "chaptercontent", "BookText", "chapter-content", "txtContent"]:
                                div = soup.find("div", id=re.compile(pid, re.I))
                                if div: break
                        if not div:
                            all_divs = soup.find_all("div")
                            if all_divs:
                                best_div = max(all_divs, key=lambda d: len(d.get_text(strip=True)))
                                if len(best_div.get_text(strip=True)) > 200: div = best_div
                        if not div:
                            first_link = None
                            for a in soup.find_all("a", href=True):
                                href = a['href']
                                if re.search(r'\d+\.html$', href) and "index" not in href:
                                    first_link = urljoin(current_url, href); break
                            if first_link:
                                self._log(f"🧭 ตรวจพบว่าเป็นหน้าสารบัญ กระโดดไปตอนแรก: {first_link}", "info")
                                current_url = first_link; success = True; break
                            else:
                                self._log(f"⚠️ [{chapter_num:04d}] ไม่พบเนื้อหาและลิงก์ตอนแรก พยายามใหม่...", "warn"); continue
                        title = soup.find("h1")
                        title_text = title.get_text(strip=True) if title else f"Chapter {chapter_num:04d}"
                        text = nv_clean_text(div, rules.get("unwanted_tags", []))
                        if len(text) < 30:
                            self._log(f"⚠️ [{chapter_num:04d}] เนื้อหาสั้นผิดปกติ พยายามใหม่...", "warn"); continue
                        safe_title = re.sub(r'[\\/*?:"<>|\r\n]', "", title_text)[:80]
                        name = (re.sub(r'[\\/*?:"<>|]', "", naming.replace("[n]", str(chapter_num)))
                                if naming.strip() else f"{chapter_num:04d}_{safe_title}")
                        with open(os.path.join(save_dir, name.strip() + ".txt"), "w", encoding="utf-8") as f:
                            f.write(f"{title_text}\n\n{text}\n\n(本集结束)")
                        self._log(f"✅ [{chapter_num:04d}] {title_text[:60]}", "ok")
                        completed += 1; chapter_num += 1
                        self.after(0, lambda c=completed: self.prog_lbl.config(text=f"โหลดสะสม: {c} ตอน"))
                        next_a = soup.find("a", string=re.compile(r'下一章|下一页|Next|หน้าถัดไป'))
                        if not next_a:
                            self._log("🏁 ไม่พบปุ่ม 'หน้าถัดไป' สิ้นสุดการทำงาน", "ok"); self._stop_event.set(); break
                        next_url = urljoin(current_url, next_a['href'])
                        is_index = any(x in next_url.lower() for x in ['mulu', 'index', 'book', 'catalog', 'info'])
                        if is_index or next_url == current_url or next_url.endswith('/'):
                            self._log("🏁 ปุ่มหน้าถัดไปชี้กลับไปหน้าสารบัญ (โหลดถึงตอนล่าสุดแล้ว)", "ok")
                            self._stop_event.set(); break
                        current_url = next_url; success = True; break
                    elif resp.status_code in (403, 429, 503):
                        self._log(f"⚠️ [{chapter_num:04d}] ติด Block ({resp.status_code}) พยายามใหม่...", "warn")
                        if resp.status_code == 403 and attempt >= 2:
                            with _nv_cf_lock:
                                nv_playwright_get_cookies(current_url, lambda m: self._log(m, "warn"), self.headless_var.get())
                            session = nv_make_session()
                    else:
                        self._log(f"⚠️ [{chapter_num:04d}] HTTP {resp.status_code} พยายามใหม่...", "warn")
                except Exception:
                    self._log(f"⚠️ [{chapter_num:04d}] Error: การเชื่อมต่อขัดข้อง พยายามใหม่...", "warn")
            if not success and not self._stop_event.is_set():
                self._log(f"❌ โหลดตอนที่ {chapter_num} ซ้ำ 5 ครั้งไม่สำเร็จ โปรแกรมหยุดอัตโนมัติ", "err")
                self._stop_event.set(); break
        self.after(0, lambda: self._on_done(completed, time.time() - t0))

    def _on_done(self, completed, elapsed):
        self._running = False
        self.start_btn.config(state="normal")
        self.stop_btn.config(state="disabled")
        self._log(f"✨ ทำงานเสร็จสิ้น โหลดสำเร็จทั้งหมด {completed} ตอน ในเวลา {elapsed:.1f} วินาที", "ok")

    def _on_stop(self):
        self._stop_event.set()
        self._log("⛔ Stopping...", "warn")
        self.stop_btn.config(state="disabled")



# ══════════════════════════════════════════════════════════════════════════════
#  MAIN LAUNCHER  — v4  (Dev-style UI + Multi-theme + 4 modules)
# ══════════════════════════════════════════════════════════════════════════════

_ML_I18N = {
    "TH": {
        "tagline":  "~/keawgood-universe  ›  select module",
        "footer":   "6 modules  ·  เปิดได้พร้อมกัน  ·  คลิกเพื่อเปิด",
        "theme_lbl":"🎨  Theme",
        "lang_lbl": "🌐  Language",
        "open":     "OPEN →",
        "apps": [
            ("novel",     "⬇", "NOVEL",     "โหลดนิยายอัตโนมัติ",     "Auto-crawl • TWKAN Hybrid Mode"),
            ("files",     "⧉", "FILES",     "จัดการไฟล์นิยาย",         "Merge · Split · Convert .txt .docx .pdf"),
            ("vocab",     "✦", "VOCAB",     "จัดการคำศัพท์นิยาย",      "De-duplicate vocab · Drag & Drop .txt"),
            ("audio",     "▶", "AUDIO",     "แปลงไฟล์เสียงเป็นวิดีโอ", "Batch .m4a/.mp3 + Image → .mp4"),
            ("clearcite", "✂", "CITE",      "ลบ Citation อัตโนมัติ",   "Remove [cite:...] tags · Batch .txt"),
            ("checker",   "✔", "CHECKER",   "เช็คตอนนิยายครบมั้ย",     "ตรวจสอบความครบถ้วนของตอน · Drag & Drop"),
        ],
    },
    "EN": {
        "tagline":  "~/keawgood-universe  ›  select module",
        "footer":   "6 modules  ·  Open simultaneously  ·  Click to launch",
        "theme_lbl":"🎨  Theme",
        "lang_lbl": "🌐  Language",
        "open":     "OPEN →",
        "apps": [
            ("novel",     "⬇", "NOVEL",     "Auto Novel Downloader",     "Auto-crawl • TWKAN Hybrid Mode"),
            ("files",     "⧉", "FILES",     "Novel File Manager",         "Merge · Split · Convert .txt .docx .pdf"),
            ("vocab",     "✦", "VOCAB",     "Vocab Optimizer",            "De-duplicate vocab · Drag & Drop .txt"),
            ("audio",     "▶", "AUDIO",     "Audio → Video Converter",    "Batch .m4a/.mp3 + Image → .mp4"),
            ("clearcite", "✂", "CITE",      "ClearCite — Remove Tags",    "Remove [cite:...] tags · Batch .txt"),
            ("checker",   "✔", "CHECKER",   "Novel Chapter Checker",      "Verify chapter completeness · Drag & Drop"),
        ],
    },
}

# Per-app accent colors (key → color pair)
_APP_ACCENTS = {
    "novel":     ("#EF4444", "#2D0C0C"),   # red
    "files":     ("#F97316", "#2D1600"),   # orange
    "vocab":     ("#22C55E", "#0A2D18"),   # green
    "audio":     ("#A855F7", "#1E0A3C"),   # purple
    "clearcite": ("#06B6D4", "#062030"),   # cyan
    "checker":   ("#F59E0B", "#2D2000"),   # amber
}

# ══════════════════════════════════════════════════════════════════════════════
#  CLEAR CITE WINDOW  — v2  (integrated into Keawgood Universe)
#  • ลบ [cite:...] [cite_start] [cite_end] ออกจากไฟล์ .txt อัตโนมัติ
#  • Drag & Drop หลายไฟล์ พร้อมกัน
#  • โหมดบันทึกแยกไฟล์ / รวมเป็นไฟล์เดียว
#  • Multi-theme aware — ใช้ระบบ Theme เดียวกับ Universe
# ══════════════════════════════════════════════════════════════════════════════

class ClearCiteWindow(ctk.CTkToplevel):
    """
    ลบ Citation Tags ออกจากไฟล์ .txt อัตโนมัติ
    รองรับ Drag & Drop, Batch Process, Merge Mode
    """

    def __init__(self, master):
        super().__init__(master)
        self.title("✦ ClearCite  ·  ลบ Citation อัตโนมัติ")
        self.geometry("720x820")
        self.minsize(580, 660)

        self._T            = get_theme()
        self.queued_files: list = []
        self.output_folder = ""
        self.merge_mode    = False
        self.files_processed = 0

        register_theme_callback(self._on_theme_change)
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        self._build()
        self._apply_theme()
        self.after(300, self._init_dnd)
        self.lift()
        self.focus_force()

    # ── Lifecycle ──────────────────────────────────────────────────────────────
    def _on_close(self):
        unregister_theme_callback(self._on_theme_change)
        self.destroy()

    def _on_theme_change(self, key):
        self._T = THEMES[key]
        self.after(50, self._apply_theme)

    # ── Regex helpers ──────────────────────────────────────────────────────────
    def _clean_content(self, content: str) -> str:
        content = re.sub(r'\[cite_start\]', '', content, flags=re.IGNORECASE)
        content = re.sub(r'\[cite_end\]',   '', content, flags=re.IGNORECASE)
        content = re.sub(r'\[cite:[^\]]*\]', '', content, flags=re.IGNORECASE | re.DOTALL)
        content = re.sub(r'[ \t]{2,}', ' ', content)
        content = re.sub(r'\n{3,}', '\n\n', content)
        return content.strip()

    def _count_citations(self, content: str) -> int:
        return len(re.findall(
            r'\[cite_start\]|\[cite_end\]|\[cite:[^\]]*\]',
            content, flags=re.IGNORECASE | re.DOTALL
        ))

    def _read_file(self, filepath: str) -> str:
        for enc in ('utf-8', 'utf-8-sig', 'cp874', 'tis-620'):
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        raise ValueError(f"อ่านไฟล์ไม่ได้: {os.path.basename(filepath)}")

    # ── File Queue ─────────────────────────────────────────────────────────────
    def _add_files(self, paths: list):
        added = 0
        for p in paths:
            p = p.strip().strip('{}').strip('"').strip("'")
            if not p or not p.lower().endswith('.txt'):
                continue
            if p not in self.queued_files:
                self.queued_files.append(p)
                added += 1
        if added:
            self._refresh_file_list()
            self._set_status(
                f"✅  เพิ่ม {added} ไฟล์  ·  รวม {len(self.queued_files)} ไฟล์",
                "success")
        else:
            self._set_status("⚠️  ไม่มีไฟล์ใหม่ (รองรับเฉพาะ .txt หรือเพิ่มซ้ำ)", "error")

    def _remove_file(self, path: str):
        if path in self.queued_files:
            self.queued_files.remove(path)
        self._refresh_file_list()
        self._set_status(f"เหลือ {len(self.queued_files)} ไฟล์", "neutral")

    def _clear_queue(self):
        self.queued_files.clear()
        self._refresh_file_list()
        self._set_status("ล้างรายการทั้งหมดแล้ว", "neutral")

    # ── Process ────────────────────────────────────────────────────────────────
    def _process_all(self):
        if not self.queued_files:
            self._set_status("⚠️  ยังไม่มีไฟล์ในคิว", "error")
            return

        out_dir = self.output_folder
        if not out_dir:
            out_dir = filedialog.askdirectory(title="เลือกโฟลเดอร์สำหรับบันทึกไฟล์")
            if not out_dir:
                self._set_status("ℹ️  ยกเลิก — ไม่ได้เลือกโฟลเดอร์", "neutral")
                return
            self.output_folder = out_dir
            self._update_folder_label()

        total_removed = 0
        errors: list = []

        if self.merge_mode:
            combined_parts = []
            for fp in self.queued_files:
                try:
                    raw = self._read_file(fp)
                    total_removed += self._count_citations(raw)
                    cleaned = self._clean_content(raw)
                    header = f"{'='*60}\n  {os.path.basename(fp)}\n{'='*60}"
                    combined_parts.append(header + "\n\n" + cleaned)
                except Exception as e:
                    errors.append(str(e))

            merged_text = "\n\n\n".join(combined_parts)
            save_name = filedialog.asksaveasfilename(
                title="บันทึกไฟล์รวม",
                defaultextension=".txt",
                filetypes=[("Text files", "*.txt")],
                initialdir=out_dir,
                initialfile="Merged_Cleaned.txt"
            )
            if save_name:
                with open(save_name, 'w', encoding='utf-8') as f:
                    f.write(merged_text)
                self.files_processed += len(self.queued_files)
                self._update_counter()
                self._set_status(
                    f"✅  รวม {len(self.queued_files)} ไฟล์สำเร็จ  ·  ลบ {total_removed} citations",
                    "success")
                self.queued_files.clear()
                self._refresh_file_list()
            else:
                self._set_status("ℹ️  ยกเลิกการบันทึก", "neutral")
        else:
            saved = 0
            for fp in self.queued_files:
                try:
                    raw = self._read_file(fp)
                    total_removed += self._count_citations(raw)
                    cleaned = self._clean_content(raw)
                    out_name = "Cleaned_" + os.path.basename(fp)
                    out_path = os.path.join(out_dir, out_name)
                    counter = 1
                    while os.path.exists(out_path):
                        base, ext = os.path.splitext(out_name)
                        out_path = os.path.join(out_dir, f"{base}_{counter}{ext}")
                        counter += 1
                    with open(out_path, 'w', encoding='utf-8') as f:
                        f.write(cleaned)
                    saved += 1
                except Exception as e:
                    errors.append(str(e))

            self.files_processed += saved
            self._update_counter()
            if errors:
                self._set_status(
                    f"⚠️  บันทึก {saved}/{len(self.queued_files)} ไฟล์  ·  มีข้อผิดพลาด {len(errors)} ไฟล์",
                    "error")
            else:
                self._set_status(
                    f"✅  บันทึก {saved} ไฟล์สำเร็จ  ·  ลบ {total_removed} citations",
                    "success")
                self.queued_files.clear()
                self._refresh_file_list()

    # ── Build UI ───────────────────────────────────────────────────────────────
    def _build(self):
        T = self._T

        # ── Header bar ────────────────────────────────────────────────────────
        self._hdr = ctk.CTkFrame(self, height=70, corner_radius=0)
        self._hdr.pack(fill="x")
        self._hdr.pack_propagate(False)

        self._hdr_title = ctk.CTkLabel(
            self._hdr, text="✦  ClearCite",
            font=mk_font(24, bold=True))
        self._hdr_title.pack(side="left", padx=24, pady=8)

        self._hdr_sub = ctk.CTkLabel(
            self._hdr, text="ลบ  [cite:...]  [cite_start]  [cite_end]  อัตโนมัติ  ·  Batch .txt",
            font=mk_font(12))
        self._hdr_sub.pack(side="left", padx=(0, 8), pady=14)

        # ── Body ──────────────────────────────────────────────────────────────
        self._body = ctk.CTkFrame(self, fg_color="transparent")
        self._body.pack(fill="both", expand=True, padx=18, pady=(12, 6))

        # ── Output folder row ─────────────────────────────────────────────────
        self._folder_card = ctk.CTkFrame(self._body, corner_radius=10, height=52)
        self._folder_card.pack(fill="x", pady=(0, 8))
        self._folder_card.pack_propagate(False)
        self._folder_card.columnconfigure(1, weight=1)

        ctk.CTkLabel(
            self._folder_card, text="📁  บันทึกไปที่:",
            font=mk_font(13, bold=True)).grid(row=0, column=0, padx=14, pady=10, sticky="w")

        self._folder_path_lbl = ctk.CTkLabel(
            self._folder_card,
            text="(ยังไม่ได้เลือก — จะถามก่อนประมวลผล)",
            font=mk_font(12), anchor="w")
        self._folder_path_lbl.grid(row=0, column=1, padx=6, pady=10, sticky="ew")

        self._btn_folder = ctk.CTkButton(
            self._folder_card, text="เลือก",
            font=mk_font(12, bold=True),
            width=80, height=32, corner_radius=8,
            command=self._choose_folder)
        self._btn_folder.grid(row=0, column=2, padx=(0, 6), pady=10)

        self._btn_folder_clear = ctk.CTkButton(
            self._folder_card, text="✕",
            font=mk_font(11, bold=True),
            width=32, height=32, corner_radius=8,
            command=self._clear_folder)
        self._btn_folder_clear.grid(row=0, column=3, padx=(0, 12), pady=10)

        # ── Merge mode toggle ─────────────────────────────────────────────────
        self._merge_card = ctk.CTkFrame(self._body, corner_radius=10, height=52)
        self._merge_card.pack(fill="x", pady=(0, 8))
        self._merge_card.pack_propagate(False)

        ctk.CTkLabel(
            self._merge_card, text="📑  โหมดบันทึก:",
            font=mk_font(13, bold=True)).pack(side="left", padx=14)

        self._btn_separate = ctk.CTkButton(
            self._merge_card, text="⬜  แยกทีละไฟล์",
            font=mk_font(12, bold=True),
            width=130, height=32, corner_radius=8,
            command=lambda: self._set_merge_mode(False))
        self._btn_separate.pack(side="left", padx=4, pady=10)

        self._btn_merge = ctk.CTkButton(
            self._merge_card, text="📋  รวมเป็นไฟล์เดียว",
            font=mk_font(12, bold=True),
            width=150, height=32, corner_radius=8,
            command=lambda: self._set_merge_mode(True))
        self._btn_merge.pack(side="left", padx=4, pady=10)

        self._merge_desc_lbl = ctk.CTkLabel(
            self._merge_card, text="→ บันทึกแยกทีละไฟล์",
            font=mk_font(11))
        self._merge_desc_lbl.pack(side="left", padx=(10, 0))

        # ── Drop zone ─────────────────────────────────────────────────────────
        self._drop_card = ctk.CTkFrame(
            self._body, corner_radius=14, height=110,
            border_width=2)
        self._drop_card.pack(fill="x", pady=(0, 10))
        self._drop_card.pack_propagate(False)

        drop_inner = ctk.CTkFrame(self._drop_card, fg_color="transparent")
        drop_inner.place(relx=0.5, rely=0.5, anchor="center")

        self._drop_icon_lbl = ctk.CTkLabel(
            drop_inner, text="📂",
            font=mk_font(28))
        self._drop_icon_lbl.pack(side="left", padx=(0, 14))

        drop_text_col = ctk.CTkFrame(drop_inner, fg_color="transparent")
        drop_text_col.pack(side="left")

        self._drop_main_lbl = ctk.CTkLabel(
            drop_text_col,
            text="ลากไฟล์ .txt หลายๆ ไฟล์มาวางที่นี่",
            font=mk_font(16, bold=True))
        self._drop_main_lbl.pack(anchor="w")

        self._drop_sub_lbl = ctk.CTkLabel(
            drop_text_col,
            text="รองรับการลากหลายไฟล์พร้อมกัน  ·  ไม่เพิ่มซ้ำ",
            font=mk_font(12))
        self._drop_sub_lbl.pack(anchor="w")

        badge_row = ctk.CTkFrame(drop_text_col, fg_color="transparent")
        badge_row.pack(anchor="w", pady=(4, 0))
        for tag in ["[cite: ...]", "[cite_start]", "[cite_end]"]:
            ctk.CTkLabel(badge_row, text=tag,
                         font=mk_font(10, bold=True, mono=True),
                         corner_radius=6, padx=8, pady=2).pack(side="left", padx=3)

        # ── File list header ───────────────────────────────────────────────────
        list_hdr = ctk.CTkFrame(self._body, fg_color="transparent", height=36)
        list_hdr.pack(fill="x", pady=(2, 4))
        list_hdr.pack_propagate(False)

        self._list_title_lbl = ctk.CTkLabel(
            list_hdr, text="รายการไฟล์ (0)",
            font=mk_font(14, bold=True))
        self._list_title_lbl.pack(side="left")

        self._btn_browse_add = ctk.CTkButton(
            list_hdr, text="+ เพิ่มไฟล์",
            font=mk_font(12, bold=True),
            width=100, height=30, corner_radius=8,
            command=self._browse_files)
        self._btn_browse_add.pack(side="right")

        self._btn_clear_all = ctk.CTkButton(
            list_hdr, text="🗑  ล้างทั้งหมด",
            font=mk_font(12),
            width=110, height=30, corner_radius=8,
            command=self._clear_queue)
        self._btn_clear_all.pack(side="right", padx=(0, 6))

        # ── Scrollable file list ───────────────────────────────────────────────
        self._list_scroll = ctk.CTkScrollableFrame(
            self._body, corner_radius=10, height=160)
        self._list_scroll.pack(fill="both", expand=True, pady=(0, 10))

        # ── Process button ─────────────────────────────────────────────────────
        self._process_btn = ctk.CTkButton(
            self._body,
            text="⚡  ประมวลผลและบันทึก",
            font=mk_font(18, bold=True),
            height=52, corner_radius=12,
            command=self._process_all)
        self._process_btn.pack(fill="x", pady=(0, 8))

        # ── Status bar ─────────────────────────────────────────────────────────
        self._status_card = ctk.CTkFrame(self._body, corner_radius=8, height=38)
        self._status_card.pack(fill="x", pady=(0, 4))
        self._status_card.pack_propagate(False)

        self._status_lbl = ctk.CTkLabel(
            self._status_card,
            text="พร้อมใช้งาน — ลากไฟล์หรือกด '+ เพิ่มไฟล์'",
            font=mk_font(12), anchor="w")
        self._status_lbl.pack(side="left", padx=14)

        self._counter_lbl = ctk.CTkLabel(
            self._status_card,
            text=f"ประมวลผลแล้ว: 0 ไฟล์",
            font=mk_font(11))
        self._counter_lbl.pack(side="right", padx=14)

        self._set_merge_mode(False)

    # ── DnD ────────────────────────────────────────────────────────────────────
    def _init_dnd(self):
        if not DND_OK:
            return
        try:
            self.drop_target_register(DND_FILES)
            self.dnd_bind('<<Drop>>', self._on_drop)
            self.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self.dnd_bind('<<DragLeave>>', self._on_drag_leave)
        except Exception:
            try:
                self.tk.call('tkdnd::drop_target', 'register', self, 'DND_Files')
                self.bind('<<Drop>>', self._on_drop)
            except Exception:
                pass

    def _on_drag_enter(self, event=None):
        T = self._T
        self._drop_card.configure(border_color=T["accent"], border_width=3)

    def _on_drag_leave(self, event=None):
        T = self._T
        self._drop_card.configure(border_color=T["bg_border"], border_width=2)

    def _on_drop(self, event):
        self._on_drag_leave()
        raw = event.data.strip() if hasattr(event, 'data') else ""
        paths = []
        if "{" in raw:
            parts = re.findall(r'\{([^}]+)\}|(\S+)', raw)
            paths = [a or b for a, b in parts]
        else:
            paths = raw.split()
        self._add_files(paths)

    # ── File list render ───────────────────────────────────────────────────────
    def _refresh_file_list(self):
        T = self._T
        for w in self._list_scroll.winfo_children():
            w.destroy()

        self._list_title_lbl.configure(text=f"รายการไฟล์ ({len(self.queued_files)})")

        if not self.queued_files:
            ctk.CTkLabel(
                self._list_scroll,
                text="ยังไม่มีไฟล์ — ลากมาวางหรือกด '+ เพิ่มไฟล์'",
                font=mk_font(12)).pack(pady=20)
            return

        for i, fp in enumerate(self.queued_files):
            row = ctk.CTkFrame(self._list_scroll, corner_radius=8, height=40)
            row.pack(fill="x", pady=2)
            row.pack_propagate(False)

            ctk.CTkLabel(
                row, text=f"{i+1:02d}",
                font=mk_font(11, bold=True, mono=True),
                width=28).pack(side="left", padx=(10, 4), pady=8)

            ctk.CTkLabel(
                row, text="📄",
                font=mk_font(13)).pack(side="left", padx=(0, 6))

            ctk.CTkLabel(
                row, text=os.path.basename(fp),
                font=mk_font(12, bold=True), anchor="w").pack(
                    side="left", fill="x", expand=True)

            dir_str = os.path.dirname(fp)
            if len(dir_str) > 40:
                dir_str = "…" + dir_str[-39:]
            ctk.CTkLabel(
                row, text=dir_str,
                font=mk_font(10), anchor="e").pack(side="left", padx=(4, 6))

            ctk.CTkButton(
                row, text="✕",
                font=mk_font(11, bold=True),
                width=28, height=28, corner_radius=6,
                command=lambda p=fp: self._remove_file(p)).pack(
                    side="right", padx=(0, 8), pady=6)

    # ── Controls ───────────────────────────────────────────────────────────────
    def _browse_files(self):
        paths = filedialog.askopenfilenames(
            title="เลือกไฟล์ .txt (Ctrl+Click เลือกหลายไฟล์)",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")])
        if paths:
            self._add_files(list(paths))

    def _choose_folder(self):
        folder = filedialog.askdirectory(title="เลือกโฟลเดอร์บันทึกไฟล์")
        if folder:
            self.output_folder = folder
            self._update_folder_label()

    def _clear_folder(self):
        self.output_folder = ""
        self._update_folder_label()

    def _update_folder_label(self):
        T = self._T
        if self.output_folder:
            short = (self.output_folder if len(self.output_folder) <= 54
                     else "…" + self.output_folder[-53:])
            self._folder_path_lbl.configure(
                text=short, text_color=T["fg_primary"])
        else:
            self._folder_path_lbl.configure(
                text="(ยังไม่ได้เลือก — จะถามก่อนประมวลผล)",
                text_color=T["fg_muted"])

    def _set_merge_mode(self, merged: bool):
        self.merge_mode = merged
        T = self._T
        if merged:
            self._btn_merge.configure(fg_color=T["accent"], text_color=T["fg_primary"])
            self._btn_separate.configure(fg_color=T["bg_input"], text_color=T["fg_muted"])
            self._merge_desc_lbl.configure(text="→ รวมทุกไฟล์ออกเป็น 1 ไฟล์")
        else:
            self._btn_separate.configure(fg_color=T["accent"], text_color=T["fg_primary"])
            self._btn_merge.configure(fg_color=T["bg_input"], text_color=T["fg_muted"])
            self._merge_desc_lbl.configure(text="→ บันทึกแยกทีละไฟล์")

    def _set_status(self, message: str, kind: str = "neutral"):
        T = self._T
        color_map = {
            "success": T["success"],
            "error":   T["error"],
            "neutral": T["fg_muted"],
        }
        color = color_map.get(kind, T["fg_muted"])
        self._status_lbl.configure(text=message, text_color=color)

    def _update_counter(self):
        self._counter_lbl.configure(
            text=f"ประมวลผลแล้ว: {self.files_processed} ไฟล์")

    # ── Theme ──────────────────────────────────────────────────────────────────
    def _apply_theme(self):
        T = self._T
        self.configure(fg_color=T["bg_root"])

        # Header
        self._hdr.configure(fg_color=T["bg_card"])
        self._hdr_title.configure(text_color=T["accent"])
        self._hdr_sub.configure(text_color=T["fg_muted"])

        # Body
        self._folder_card.configure(fg_color=T["bg_card"])
        self._folder_path_lbl.configure(text_color=T["fg_muted"])
        self._btn_folder.configure(
            fg_color=T["accent"], hover_color=T["accent2"],
            text_color=T["fg_primary"])
        self._btn_folder_clear.configure(
            fg_color=T["error"], hover_color="#B91C1C",
            text_color="#FFFFFF")

        self._merge_card.configure(fg_color=T["bg_card"])
        self._merge_desc_lbl.configure(text_color=T["fg_muted"])
        self._set_merge_mode(self.merge_mode)

        # Drop zone
        self._drop_card.configure(
            fg_color=T["bg_card"], border_color=T["bg_border"])
        self._drop_main_lbl.configure(text_color=T["fg_primary"])
        self._drop_sub_lbl.configure(text_color=T["fg_muted"])
        self._drop_icon_lbl.configure(text_color=T["fg_muted"])

        # List
        self._list_title_lbl.configure(text_color=T["fg_primary"])
        self._btn_browse_add.configure(
            fg_color=T["accent"], hover_color=T["accent2"],
            text_color=T["fg_primary"])
        self._btn_clear_all.configure(
            fg_color=T["error"], hover_color="#B91C1C",
            text_color="#FFFFFF")
        self._list_scroll.configure(fg_color=T["bg_input"])

        # Process btn
        self._process_btn.configure(
            fg_color=T["accent"], hover_color=T["accent2"],
            text_color=T["fg_primary"])

        # Status
        self._status_card.configure(fg_color=T["bg_card"])
        self._status_lbl.configure(text_color=T["fg_muted"])
        self._counter_lbl.configure(text_color=T["fg_muted"])

        # Update folder label colors
        self._update_folder_label()
        # Re-render file rows with new theme
        self._refresh_file_list()


# ──────────────────────────────────────────────────────────────────────────────


# ══════════════════════════════════════════════════════════════════════════════
#  NOVEL CHECKER WINDOW  — v2  (merged from novel_checker.py + upgraded UX)
#  • เช็คความครบถ้วนของตอนนิยายจากเนื้อหาไฟล์
#  • รองรับ Drag & Drop หลายไฟล์พร้อมกัน (โฟลเดอร์ได้ด้วย)
#  • รองรับ .txt .docx .md .pdf
#  • Multi-theme + Multi-language aware (ใช้ระบบเดียวกับ Universe)
#  • ลบโค้ดซ้ำจาก novel_checker.py และปรับรวมกับ bk_read_text()
# ══════════════════════════════════════════════════════════════════════════════

_NC_I18N = {
    "TH": {
        "title":        "✦  Novel Checker  ·  เช็คตอนนิยาย",
        "sub":          "ตรวจสอบความครบถ้วนของตอนจากเนื้อหาไฟล์",
        "sec_range":    "📐  ช่วงตอนที่ต้องการตรวจ",
        "lbl_from":     "ตอนเริ่มต้น:",
        "lbl_to":       "ตอนสิ้นสุด:",
        "sec_format":   "🔤  รูปแบบคำนำหน้าตอนในเนื้อหา",
        "sec_drop":     "📂  ลากไฟล์หรือโฟลเดอร์มาวางที่นี่",
        "drop_hint":    "รองรับ .txt · .docx · .md · .pdf  |  ลากทีละกี่ไฟล์ก็ได้  |  รองรับโฟลเดอร์",
        "sec_log":      "📋  ผลการสแกน",
        "btn_clear":    "🗑  ล้างหน้าจอ",
        "btn_scan":     "🔍  เริ่มสแกน",
        "err_input":    "⚠️  กรุณากรอก 'ตอนเริ่มต้น' และ 'ตอนสิ้นสุด' เป็นตัวเลข",
        "err_range":    "⚠️  'ตอนเริ่มต้น' ต้องน้อยกว่าหรือเท่ากับ 'ตอนสิ้นสุด'",
        "err_nofiles":  "⚠️  ไม่พบไฟล์ที่รองรับ (.txt .docx .md .pdf)",
        "msg_found":    "📁  พบไฟล์ที่รองรับ {n} ไฟล์  —  กำลังสแกน...",
        "msg_scanning": "🔍  กำลังสแกน: {f}",
        "msg_skip":     "   ⚠️  ข้ามไฟล์ {f}  (อ่านไม่ได้หรือว่างเปล่า)",
        "msg_ok":       "✅  ครบถ้วน!  ทุกตอนตั้งแต่ {s} ถึง {e} มีอยู่ครบ",
        "msg_missing":  "❌  ขาดหายไป {n} ตอน:  {lst}",
        "formats_th":   ["ตอนที่ [n]", "Chapter [n]", "第[n]章 (ภาษาจีน)", "Auto (ค้นหาอัตโนมัติ)"],
        "formats_en":   ["ตอนที่ [n]", "Chapter [n]", "第[n]章 (Chinese)", "Auto (Any prefix)"],
        "ready":        "✨  พร้อมใช้งาน!  ระบุช่วงตอน เลือกรูปแบบ แล้วลากไฟล์ลงมาได้เลย",
        "reset_ok":     "🔄  รีเซ็ตเรียบร้อย  —  พร้อมสแกนใหม่",
    },
    "EN": {
        "title":        "✦  Novel Checker  ·  Chapter Verifier",
        "sub":          "Verify chapter completeness from file contents",
        "sec_range":    "📐  Chapter Range",
        "lbl_from":     "Start chapter:",
        "lbl_to":       "End chapter:",
        "sec_format":   "🔤  Chapter Prefix Format",
        "sec_drop":     "📂  Drag & Drop files or folders here",
        "drop_hint":    "Supports .txt · .docx · .md · .pdf  |  Multiple files  |  Folders OK",
        "sec_log":      "📋  Scan Results",
        "btn_clear":    "🗑  Clear",
        "btn_scan":     "🔍  Scan",
        "err_input":    "⚠️  Please enter valid numbers for Start and End.",
        "err_range":    "⚠️  Start must be ≤ End.",
        "err_nofiles":  "⚠️  No supported files found (.txt .docx .md .pdf)",
        "msg_found":    "📁  Found {n} supported file(s)  —  Scanning...",
        "msg_scanning": "🔍  Scanning: {f}",
        "msg_skip":     "   ⚠️  Skipped {f}  (unreadable or empty)",
        "msg_ok":       "✅  Complete!  All chapters from {s} to {e} are present.",
        "msg_missing":  "❌  Missing {n} chapter(s):  {lst}",
        "formats_th":   ["ตอนที่ [n]", "Chapter [n]", "第[n]章 (ภาษาจีน)", "Auto (ค้นหาอัตโนมัติ)"],
        "formats_en":   ["ตอนที่ [n]", "Chapter [n]", "第[n]章 (Chinese)", "Auto (Any prefix)"],
        "ready":        "✨  Ready! Set chapter range, choose format, then drop files.",
        "reset_ok":     "🔄  Reset complete — ready for a new scan.",
    },
}


class NovelCheckerWindow(ctk.CTkToplevel):
    """
    เช็คความครบถ้วนของตอนนิยายจากเนื้อหาไฟล์
    รวมจาก novel_checker.py และปรับปรุง UX/UI ให้สอดคล้องกับ Universe v5
    - ใช้ bk_read_text() ร่วมกัน (ไม่ซ้ำโค้ด)
    - ใช้ get_theme() / get_lang() ของ Universe
    - ลบ TkDnDApp subclass ออก เพราะ root ถูก require แล้วจาก MainLauncher
    """

    def __init__(self, master):
        super().__init__(master)
        self._lang = get_lang()
        self._T    = get_theme()
        self._scanning = False

        register_theme_callback(self._on_theme_change)
        register_lang_callback(self._on_lang_change)
        self.protocol("WM_DELETE_WINDOW", self._on_close)

        self._setup_window()
        self._build()
        self._apply_theme()
        self.after(300, self._init_dnd)
        self.lift()
        self.focus_force()

    # ── Lifecycle ──────────────────────────────────────────────────────────────
    def _on_close(self):
        unregister_theme_callback(self._on_theme_change)
        unregister_lang_callback(self._on_lang_change)
        self.destroy()

    def _on_theme_change(self, key):
        self._T = THEMES[key]
        self.after(50, self._apply_theme)

    def _on_lang_change(self, lang):
        self._lang = lang
        self.after(50, self._refresh_lang)

    def _L(self):
        return _NC_I18N.get(self._lang, _NC_I18N["TH"])

    # ── Window setup ───────────────────────────────────────────────────────────
    def _setup_window(self):
        L = self._L()
        self.title(L["title"])
        self.geometry("820x780")
        self.minsize(680, 640)

    # ── Build UI ───────────────────────────────────────────────────────────────
    def _build(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(4, weight=1)

        self._build_header()
        self._build_range_section()
        self._build_format_section()
        self._build_drop_zone()
        self._build_log_section()

    # ── Header bar ─────────────────────────────────────────────────────────────
    def _build_header(self):
        L = self._L()
        self._hdr = ctk.CTkFrame(self, height=68, corner_radius=0)
        self._hdr.grid(row=0, column=0, sticky="ew")
        self._hdr.grid_propagate(False)
        self._hdr.columnconfigure(1, weight=1)

        # Accent bar
        self._hdr_accent = ctk.CTkFrame(self._hdr, height=3, corner_radius=0)
        self._hdr_accent.grid(row=0, column=0, columnspan=3, sticky="ew")

        self._hdr_title = ctk.CTkLabel(
            self._hdr, text=L["title"],
            font=mk_font(22, bold=True))
        self._hdr_title.grid(row=1, column=0, padx=22, pady=(8, 8), sticky="w")

        self._hdr_sub = ctk.CTkLabel(
            self._hdr, text=L["sub"],
            font=mk_font(11))
        self._hdr_sub.grid(row=1, column=1, padx=4, sticky="w")

    # ── Range section ──────────────────────────────────────────────────────────
    def _build_range_section(self):
        L = self._L()
        self._range_card = ctk.CTkFrame(self, corner_radius=12)
        self._range_card.grid(row=1, column=0, padx=18, pady=(12, 6), sticky="ew")
        self._range_card.columnconfigure((1, 3), weight=1)

        self._lbl_sec_range = ctk.CTkLabel(
            self._range_card, text=L["sec_range"],
            font=mk_font(14, bold=True))
        self._lbl_sec_range.grid(row=0, column=0, columnspan=5, padx=16, pady=(12, 8), sticky="w")

        self._lbl_from = ctk.CTkLabel(
            self._range_card, text=L["lbl_from"],
            font=mk_font(13))
        self._lbl_from.grid(row=1, column=0, padx=(16, 6), pady=(0, 14), sticky="w")

        self._entry_from = ctk.CTkEntry(
            self._range_card, placeholder_text="1",
            width=100, height=36, font=mk_font(15, bold=True),
            justify="center", corner_radius=8)
        self._entry_from.grid(row=1, column=1, padx=(0, 20), pady=(0, 14), sticky="w")

        self._lbl_to = ctk.CTkLabel(
            self._range_card, text=L["lbl_to"],
            font=mk_font(13))
        self._lbl_to.grid(row=1, column=2, padx=(0, 6), pady=(0, 14), sticky="w")

        self._entry_to = ctk.CTkEntry(
            self._range_card, placeholder_text="500",
            width=100, height=36, font=mk_font(15, bold=True),
            justify="center", corner_radius=8)
        self._entry_to.grid(row=1, column=3, padx=(0, 20), pady=(0, 14), sticky="w")

        self._btn_scan = ctk.CTkButton(
            self._range_card, text=L["btn_scan"],
            font=mk_font(14, bold=True),
            height=36, width=130, corner_radius=8,
            command=self._run_from_entries)
        self._btn_scan.grid(row=1, column=4, padx=(0, 16), pady=(0, 14))

    # ── Format section ─────────────────────────────────────────────────────────
    def _build_format_section(self):
        L = self._L()
        self._fmt_card = ctk.CTkFrame(self, corner_radius=12)
        self._fmt_card.grid(row=2, column=0, padx=18, pady=(0, 6), sticky="ew")
        self._fmt_card.columnconfigure(1, weight=1)

        self._lbl_sec_fmt = ctk.CTkLabel(
            self._fmt_card, text=L["sec_format"],
            font=mk_font(14, bold=True))
        self._lbl_sec_fmt.grid(row=0, column=0, padx=16, pady=(12, 8), sticky="w")

        formats = L["formats_th"] if self._lang == "TH" else L["formats_en"]
        self._fmt_var = ctk.StringVar(value=formats[0])
        self._fmt_menu = ctk.CTkOptionMenu(
            self._fmt_card, variable=self._fmt_var,
            values=formats,
            font=mk_font(13), width=260, height=34, corner_radius=8)
        self._fmt_menu.grid(row=0, column=1, padx=(0, 16), pady=(12, 12), sticky="w")

    # ── Drop zone ──────────────────────────────────────────────────────────────
    def _build_drop_zone(self):
        L = self._L()
        self._drop_card = ctk.CTkFrame(
            self, corner_radius=14, height=120, border_width=2)
        self._drop_card.grid(row=3, column=0, padx=18, pady=(0, 6), sticky="ew")
        self._drop_card.grid_propagate(False)

        inner = ctk.CTkFrame(self._drop_card, fg_color="transparent")
        inner.place(relx=0.5, rely=0.5, anchor="center")

        self._drop_icon = ctk.CTkLabel(
            inner, text="📂", font=mk_font(32))
        self._drop_icon.pack(side="left", padx=(0, 18))

        txt_col = ctk.CTkFrame(inner, fg_color="transparent")
        txt_col.pack(side="left")

        self._drop_main = ctk.CTkLabel(
            txt_col, text=L["sec_drop"],
            font=mk_font(17, bold=True))
        self._drop_main.pack(anchor="w")

        self._drop_hint = ctk.CTkLabel(
            txt_col, text=L["drop_hint"],
            font=mk_font(11))
        self._drop_hint.pack(anchor="w", pady=(3, 0))

    # ── Log section ────────────────────────────────────────────────────────────
    def _build_log_section(self):
        L = self._L()
        log_outer = ctk.CTkFrame(self, corner_radius=12)
        log_outer.grid(row=4, column=0, padx=18, pady=(0, 18), sticky="nsew")
        log_outer.grid_rowconfigure(1, weight=1)
        log_outer.grid_columnconfigure(0, weight=1)

        hdr_row = ctk.CTkFrame(log_outer, fg_color="transparent", height=38)
        hdr_row.grid(row=0, column=0, sticky="ew", padx=14, pady=(10, 4))
        hdr_row.grid_propagate(False)
        hdr_row.columnconfigure(0, weight=1)

        self._lbl_sec_log = ctk.CTkLabel(
            hdr_row, text=L["sec_log"],
            font=mk_font(14, bold=True))
        self._lbl_sec_log.grid(row=0, column=0, sticky="w")

        self._btn_clear = ctk.CTkButton(
            hdr_row, text=L["btn_clear"],
            font=mk_font(12, bold=True),
            height=30, width=120, corner_radius=8,
            command=self._clear_log)
        self._btn_clear.grid(row=0, column=1, sticky="e")

        self._log_box = ctk.CTkTextbox(
            log_outer, font=mk_font(13, mono=True),
            corner_radius=10, wrap="word", state="disabled")
        self._log_box.grid(row=1, column=0, padx=10, pady=(0, 10), sticky="nsew")

        self._write_log(L["ready"])

    # ── DnD ────────────────────────────────────────────────────────────────────
    def _init_dnd(self):
        if not DND_OK:
            return
        try:
            self.drop_target_register(DND_FILES)
            self.dnd_bind('<<Drop>>', self._on_drop)
            self.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self.dnd_bind('<<DragLeave>>', self._on_drag_leave)
        except Exception:
            try:
                self.tk.call('tkdnd::drop_target', 'register', self, 'DND_Files')
                self.bind('<<Drop>>', self._on_drop)
            except Exception:
                pass

    def _on_drag_enter(self, event=None):
        self._drop_card.configure(border_width=3)

    def _on_drag_leave(self, event=None):
        self._drop_card.configure(border_width=2)

    def _on_drop(self, event):
        self._on_drag_leave()
        raw = event.data if hasattr(event, 'data') else ""
        raw = raw.strip()
        if "{" in raw:
            parts = re.findall(r'\{([^}]+)\}|(\S+)', raw)
            paths = [a or b for a, b in parts]
        else:
            paths = self.tk.splitlist(raw)
        self._run_scan(list(paths))

    # ── Core Logic ─────────────────────────────────────────────────────────────
    def _get_regex(self) -> str:
        choice = self._fmt_var.get()
        if "ตอนที่" in choice:
            return r'ตอนที่\s*(\d+)'
        elif "Chapter" in choice:
            return r'(?:chapter|ch|ch\.)\s*(\d+)'
        elif "第" in choice:
            return r'第\s*(\d+)\s*[章ตอน]'
        else:
            return r'(?:ตอนที่|chapter|ch\.|第)\s*(\d+)'

    def _collect_files(self, paths: list) -> list:
        valid_exts = ('.txt', '.docx', '.md', '.pdf')
        result = []
        for p in paths:
            p = p.strip().strip('{}').strip('"').strip("'")
            if not p:
                continue
            if os.path.isdir(p):
                for root, _, files in os.walk(p):
                    for f in files:
                        if f.lower().endswith(valid_exts):
                            result.append(os.path.join(root, f))
            elif os.path.isfile(p) and p.lower().endswith(valid_exts):
                result.append(p)
        return result

    def _run_from_entries(self):
        """Called by the Scan button — builds a synthetic file-drop from drop zone.
           Requires files to already have been dropped. Uses last-drop paths cache."""
        # If no files were dropped, prompt via dialog
        paths = filedialog.askopenfilenames(
            title="เลือกไฟล์ที่ต้องการสแกน (Ctrl+Click เลือกหลายไฟล์)",
            filetypes=[
                ("Supported files", "*.txt *.docx *.md *.pdf"),
                ("Text", "*.txt"), ("Word", "*.docx"),
                ("Markdown", "*.md"), ("PDF", "*.pdf"),
                ("All", "*.*"),
            ])
        if paths:
            self._run_scan(list(paths))

    def _run_scan(self, raw_paths: list):
        if self._scanning:
            return
        L = self._L()
        # Validate range
        try:
            start = int(self._entry_from.get().strip())
            end   = int(self._entry_to.get().strip())
        except ValueError:
            self._write_log(L["err_input"])
            return
        if start > end:
            self._write_log(L["err_range"])
            return

        files = self._collect_files(raw_paths)
        if not files:
            self._write_log(L["err_nofiles"])
            return

        self._scanning = True
        self._write_log("─" * 58)
        self._write_log(L["msg_found"].format(n=len(files)))

        pattern = re.compile(self._get_regex(), re.IGNORECASE)
        found_chapters: set = set()

        for fp in files:
            fname = os.path.basename(fp)
            self._write_log(L["msg_scanning"].format(f=fname))
            try:
                content = bk_read_text(fp)
                for m in pattern.finditer(content):
                    found_chapters.add(int(m.group(1)))
            except Exception:
                self._write_log(L["msg_skip"].format(f=fname))

        missing = [i for i in range(start, end + 1) if i not in found_chapters]

        if not missing:
            self._write_log(L["msg_ok"].format(s=start, e=end))
        else:
            lst_str = ", ".join(map(str, missing))
            self._write_log(L["msg_missing"].format(n=len(missing), lst=lst_str))

        self._write_log("─" * 58)
        self._scanning = False

    # ── Log helpers ────────────────────────────────────────────────────────────
    def _write_log(self, text: str):
        self._log_box.configure(state="normal")
        self._log_box.insert("end", text + "\n")
        self._log_box.see("end")
        self._log_box.configure(state="disabled")
        self.update_idletasks()

    def _clear_log(self):
        L = self._L()
        self._log_box.configure(state="normal")
        self._log_box.delete("1.0", "end")
        self._log_box.configure(state="disabled")
        self._entry_from.delete(0, "end")
        self._entry_to.delete(0, "end")
        self._write_log(L["reset_ok"])

    # ── Lang refresh ───────────────────────────────────────────────────────────
    def _refresh_lang(self):
        L = self._L()
        try:
            self.title(L["title"])
            self._hdr_title.configure(text=L["title"])
            self._hdr_sub.configure(text=L["sub"])
            self._lbl_sec_range.configure(text=L["sec_range"])
            self._lbl_from.configure(text=L["lbl_from"])
            self._lbl_to.configure(text=L["lbl_to"])
            self._btn_scan.configure(text=L["btn_scan"])
            self._lbl_sec_fmt.configure(text=L["sec_format"])
            formats = L["formats_th"] if self._lang == "TH" else L["formats_en"]
            self._fmt_menu.configure(values=formats)
            if self._fmt_var.get() not in formats:
                self._fmt_var.set(formats[0])
            self._drop_main.configure(text=L["sec_drop"])
            self._drop_hint.configure(text=L["drop_hint"])
            self._lbl_sec_log.configure(text=L["sec_log"])
            self._btn_clear.configure(text=L["btn_clear"])
        except Exception:
            pass

    # ── Theme apply ────────────────────────────────────────────────────────────
    def _apply_theme(self):
        T = self._T
        self.configure(fg_color=T["bg_root"])

        self._hdr.configure(fg_color=T["bg_card"])
        try:
            self._hdr_accent.configure(fg_color=T["accent"])
        except Exception:
            pass
        self._hdr_title.configure(text_color=T["accent"])
        self._hdr_sub.configure(text_color=T["fg_muted"])

        for card in [self._range_card, self._fmt_card]:
            card.configure(fg_color=T["bg_card"])

        for lbl in [self._lbl_sec_range, self._lbl_from, self._lbl_to, self._lbl_sec_fmt]:
            lbl.configure(text_color=T["fg_primary"])

        for entry in [self._entry_from, self._entry_to]:
            entry.configure(
                fg_color=T["bg_input"],
                border_color=T["bg_border"],
                text_color=T["fg_body"])

        self._btn_scan.configure(
            fg_color=T["accent"], hover_color=T["accent2"],
            text_color=T["fg_primary"])

        self._fmt_menu.configure(
            fg_color=T["bg_input"],
            button_color=T["accent"],
            button_hover_color=T["accent2"],
            text_color=T["fg_body"])

        self._drop_card.configure(
            fg_color=T["bg_card"], border_color=T["bg_border"])
        self._drop_main.configure(text_color=T["fg_primary"])
        self._drop_hint.configure(text_color=T["fg_muted"])
        self._drop_icon.configure(text_color=T["fg_muted"])

        self._log_box.configure(
            fg_color=T["console_bg"],
            text_color=T["console_fg"],
            border_color=T["bg_border"])

        self._btn_clear.configure(
            fg_color=T["error"], hover_color="#B91C1C",
            text_color="#FFFFFF")


# ──────────────────────────────────────────────────────────────────────────────

_APP_FACTORIES = {
    "novel":     lambda hub: NovelByKeawgoodWindow(hub),
    "files":     lambda hub: ByKeawgoodWindow(hub),
    "vocab":     lambda hub: VocabOptimizerWindow(hub),
    "audio":     lambda hub: AudioByKeawgoodWindow(hub),
    "clearcite": lambda hub: ClearCiteWindow(hub),
    "checker":   lambda hub: NovelCheckerWindow(hub),
}


class MainLauncher(ctk.CTk):
    """
    Keawgood Universe — Main Launcher
    Design: VS Code sidebar · Linear cards · Vercel dashboard
    """

    def __init__(self):
        super().__init__()
        self._lang = get_lang()
        self._T    = get_theme()
        self._wins: dict = {}

        # Init DnD on the root window (required once)
        if DND_OK:
            self.after(100, lambda: _dnd_require_root(self))

        register_theme_callback(self._on_theme_change)
        register_lang_callback(self._on_lang_change)

        self._setup_window()
        self._build()
        self._apply_theme()

    # ── Window setup ───────────────────────────────────────────────────────────
    def _setup_window(self):
        self.title("Keawgood Universe  v5.0")
        self.geometry("600x980")
        self.minsize(540, 840)
        self.resizable(True, True)
        ctk.set_appearance_mode(self._T["ctk_mode"])

    # ── Theme / Lang callbacks ─────────────────────────────────────────────────
    def _on_theme_change(self, key):
        self._T = THEMES[key]
        self.after(50, self._apply_theme)

    def _on_lang_change(self, lang):
        self._lang = lang
        self.after(50, self._rebuild_cards)

    def _L(self): return _ML_I18N.get(self._lang, _ML_I18N["TH"])

    # ── Build ──────────────────────────────────────────────────────────────────
    def _build(self):
        self.grid_rowconfigure(2, weight=1)
        self.grid_columnconfigure(0, weight=1)
        self._build_topbar()
        self._build_hero()
        self._build_cards_area()
        self._build_statusbar()

    # ── Top bar (VSCode style) ─────────────────────────────────────────────────
    def _build_topbar(self):
        self._topbar = tk.Frame(self, height=36)
        self._topbar.grid(row=0, column=0, sticky="ew")
        self._topbar.grid_propagate(False)
        self._topbar.columnconfigure(3, weight=1)

        # Traffic-light dots
        dots = tk.Frame(self._topbar)
        dots.grid(row=0, column=0, padx=12, pady=9, sticky="w")
        for col in ("#FF5F57", "#FFBD2E", "#28C940"):
            c = tk.Canvas(dots, width=12, height=12,
                          highlightthickness=0)
            c.pack(side="left", padx=2)
            c.create_oval(1, 1, 11, 11, fill=col, outline="")
        self._dots_canvases = [dots.winfo_children()[i] for i in range(3)] \
            if dots.winfo_children() else []

        # App breadcrumb
        self._lbl_crumb = tk.Label(
            self._topbar, text="keawgood-universe",
            font=tk_font(10, mono=True))
        self._lbl_crumb.grid(row=0, column=1, padx=4, pady=9, sticky="w")

        # Spacer
        tk.Frame(self._topbar).grid(row=0, column=3)

        # Theme selector
        self._theme_var = tk.StringVar(value=_ACTIVE_THEME_KEY)
        self._theme_labels = [THEMES[k]["label"] for k in THEME_KEYS]
        self._theme_cb = ttk.Combobox(
            self._topbar,
            textvariable=self._theme_var,
            values=self._theme_labels,
            width=14, state="readonly",
            font=tk_font(9, mono=True))
        self._theme_cb.grid(row=0, column=4, padx=4, pady=6)
        self._theme_cb.bind("<<ComboboxSelected>>", self._on_theme_select)

        # Language button
        self._lbl_langbtn = tk.Button(
            self._topbar, text="🌐  TH/EN",
            font=tk_font(9, mono=True),
            relief="flat", bd=0, cursor="hand2",
            padx=8, pady=2,
            command=self._toggle_lang)
        self._lbl_langbtn.grid(row=0, column=5, padx=(0, 10), pady=6)

    # ── Hero ───────────────────────────────────────────────────────────────────
    def _build_hero(self):
        self._hero = tk.Frame(self, height=118)
        self._hero.grid(row=1, column=0, sticky="ew")
        self._hero.grid_propagate(False)

        # Gradient accent strip (drawn on canvas)
        self._strip = tk.Canvas(self._hero, height=3, highlightthickness=0)
        self._strip.pack(fill="x")
        self._strip.bind("<Configure>", self._draw_strip)

        inner = tk.Frame(self._hero)
        inner.pack(expand=True)

        # Version badge
        self._badge_frame = tk.Frame(inner, padx=6, pady=1)
        self._badge_frame.pack(pady=(8, 3))
        self._badge_lbl = tk.Label(
            self._badge_frame,
            text="  v5.0  ·  6 Modules  ",
            font=tk_font(9, mono=True))
        self._badge_lbl.pack()

        # Title
        self._title_lbl = tk.Label(
            inner, text="Keawgood Universe",
            font=tk_font(24, bold=True, mono=True))
        self._title_lbl.pack()

        # Tagline
        self._tag_lbl = tk.Label(
            inner, text=self._L()["tagline"],
            font=tk_font(10, mono=True))
        self._tag_lbl.pack(pady=(1, 8))

    def _draw_strip(self, event=None):
        try:
            w = self._strip.winfo_width()
            if w < 4: return
            self._strip.delete("all")
            stops = [
                (0.00, "#6366F1"),
                (0.33, "#A855F7"),
                (0.66, "#EC4899"),
                (1.00, "#F97316"),
            ]
            step = max(1, w // 120)
            for i in range(0, w, step):
                t = i / w
                # find segment
                c = "#6366F1"
                for j in range(len(stops) - 1):
                    if stops[j][0] <= t <= stops[j+1][0]:
                        seg_t = (t - stops[j][0]) / (stops[j+1][0] - stops[j][0] + 1e-9)
                        c = self._lerp(stops[j][1], stops[j+1][1], seg_t)
                        break
                self._strip.create_rectangle(i, 0, i + step, 4, fill=c, outline="")
        except Exception:
            pass

    @staticmethod
    def _lerp(c1, c2, t):
        r1,g1,b1 = int(c1[1:3],16),int(c1[3:5],16),int(c1[5:7],16)
        r2,g2,b2 = int(c2[1:3],16),int(c2[3:5],16),int(c2[5:7],16)
        return "#{:02x}{:02x}{:02x}".format(
            int(r1+(r2-r1)*t),int(g1+(g2-g1)*t),int(b1+(b2-b1)*t))

    # ── Cards area ─────────────────────────────────────────────────────────────
    def _build_cards_area(self):
        # Outer container with scrollbar (prevents clipping on small screens)
        self._scroll = ctk.CTkScrollableFrame(
            self, corner_radius=0,
            scrollbar_button_color="#252542",
            scrollbar_button_hover_color="#6366F1")
        self._scroll.grid(row=2, column=0, sticky="nsew")
        self._scroll.columnconfigure(0, weight=1)
        self._cards_widgets = []
        self._build_cards_inner()

    def _build_cards_inner(self):
        L = self._L()
        # Section comment label
        self._sect_lbl = tk.Label(
            self._scroll,
            text="  // modules — click to launch",
            font=tk_font(10, mono=True),
            anchor="w")
        self._sect_lbl.pack(fill="x", padx=18, pady=(12, 6))

        self._card_frames = []
        for app_data in L["apps"]:
            key, icon, tag, name, desc = app_data
            ac, dim = _APP_ACCENTS[key]
            frame = self._make_card(key, icon, tag, name, desc, ac, dim)
            self._card_frames.append((frame, ac, dim, key, icon, tag, name, desc))

        tk.Frame(self._scroll, height=10).pack()

    def _make_card(self, key, icon, tag, name, desc, ac, dim):
        T = self._T
        card = tk.Frame(
            self._scroll,
            highlightthickness=1,
            cursor="hand2")
        card.pack(fill="x", padx=18, pady=5)

        # Left color bar
        bar = tk.Frame(card, width=5)
        bar.pack(side="left", fill="y")

        # Icon badge
        ic = tk.Canvas(card, width=54, height=54, highlightthickness=0)
        ic.pack(side="left", padx=(12, 12), pady=12)
        ic.create_text(27, 27, text=icon, font=tk_font(20, bold=True, mono=True))

        # Text
        txt = tk.Frame(card)
        txt.pack(side="left", fill="both", expand=True, pady=12)

        top = tk.Frame(txt)
        top.pack(anchor="w")

        tag_lbl = tk.Label(top, text=f" {tag} ",
                            font=tk_font(8, bold=True, mono=True),
                            padx=4, pady=1)
        tag_lbl.pack(side="left")

        name_lbl = tk.Label(top, text=f"  {name}",
                             font=tk_font(16, bold=True))
        name_lbl.pack(side="left")

        desc_lbl = tk.Label(txt, text=desc,
                             font=tk_font(9, mono=True),
                             anchor="w")
        desc_lbl.pack(anchor="w", pady=(2, 0))

        # Launch button
        open_btn = tk.Button(
            card, text=self._L()["open"],
            font=tk_font(9, bold=True, mono=True),
            relief="flat", bd=0, cursor="hand2",
            padx=12, pady=6,
            command=lambda k=key: self._launch(k))
        open_btn.pack(side="right", padx=14, pady=16)

        # store refs for theme / hover
        _refs = dict(card=card, bar=bar, ic=ic, txt=txt, top=top,
                     tag_lbl=tag_lbl, name_lbl=name_lbl, desc_lbl=desc_lbl,
                     open_btn=open_btn, ac=ac, dim=dim)

        def _theme_card(r=_refs):
            T = self._T
            r["card"].configure(bg=T["bg_card"], highlightbackground=T["bg_border"])
            r["bar"].configure(bg=r["ac"])
            r["ic"].configure(bg=r["dim"])
            r["ic"].itemconfig("all", fill=r["ac"])
            r["txt"].configure(bg=T["bg_card"])
            r["top"].configure(bg=T["bg_card"])
            r["tag_lbl"].configure(bg=r["dim"], fg=r["ac"])
            r["name_lbl"].configure(bg=T["bg_card"], fg=T["fg_primary"])
            r["desc_lbl"].configure(bg=T["bg_card"], fg=T["fg_muted"])
            r["open_btn"].configure(bg=r["dim"], fg=r["ac"],
                                    activebackground=r["ac"],
                                    activeforeground=T["fg_primary"])
        _refs["_theme_card"] = _theme_card
        _theme_card()

        def _enter(e, r=_refs):
            T = self._T
            r["card"].configure(bg=T["bg_card_hv"], highlightbackground=r["ac"])
            for w in [r["txt"], r["top"], r["name_lbl"], r["desc_lbl"]]:
                try: w.configure(bg=T["bg_card_hv"])
                except Exception: pass
            r["open_btn"].configure(bg=r["ac"], fg=T["fg_primary"])

        def _leave(e, r=_refs):
            _theme_card()

        for w in [card, ic, txt, top, name_lbl, desc_lbl, tag_lbl]:
            w.bind("<Enter>", _enter)
            w.bind("<Leave>", _leave)
            w.bind("<Button-1>", lambda e, k=key: self._launch(k))

        # Store theme refresh function
        card._theme_card = _theme_card
        return card

    def _rebuild_cards(self):
        """Re-render cards when language changes."""
        # Destroy old
        try:
            self._sect_lbl.destroy()
            for frame, *_ in self._card_frames:
                frame.destroy()
        except Exception:
            pass
        self._build_cards_inner()
        self._apply_theme()

    # ── Status bar ─────────────────────────────────────────────────────────────
    def _build_statusbar(self):
        self._statusbar = tk.Frame(self, height=26)
        self._statusbar.grid(row=3, column=0, sticky="ew")
        self._statusbar.grid_propagate(False)

        self._sb_left = tk.Label(
            self._statusbar, text="⎇  main",
            font=tk_font(9, mono=True), anchor="w")
        self._sb_left.pack(side="left", padx=12)

        self._sb_mid = tk.Label(
            self._statusbar, text=self._L()["footer"],
            font=tk_font(9, mono=True))
        self._sb_mid.pack(side="left")

        self._sb_right = tk.Label(
            self._statusbar, text="Python  ·  v5.0  ✓",
            font=tk_font(9, mono=True))
        self._sb_right.pack(side="right", padx=12)

    # ── Theme ──────────────────────────────────────────────────────────────────
    def _on_theme_select(self, event=None):
        label = self._theme_var.get()
        for k, v in THEMES.items():
            if v["label"] == label:
                set_theme(k)
                break

    def _apply_theme(self):
        T = self._T
        self.configure(fg_color=T["bg_root"])

        # Topbar
        self._topbar.configure(bg=T["bg_sidebar"])
        self._lbl_crumb.configure(bg=T["bg_sidebar"], fg=T["fg_muted"])
        self._lbl_langbtn.configure(
            bg=T["bg_input"], fg=T["fg_muted"],
            activebackground=T["accent"], activeforeground=T["fg_primary"])
        # Combobox style
        try:
            style = ttk.Style()
            style.theme_use("clam")
            style.configure("TCombobox",
                             fieldbackground=T["bg_input"],
                             background=T["bg_input"],
                             foreground=T["fg_muted"],
                             selectbackground=T["accent"],
                             selectforeground=T["fg_primary"],
                             bordercolor=T["bg_border"],
                             arrowcolor=T["fg_muted"])
        except Exception:
            pass

        # Hero
        self._hero.configure(bg=T["bg_card"])
        self._strip.configure(bg=T["bg_card"])
        self._badge_frame.configure(bg=T["accent"].replace("F1","40") if "F1" in T["accent"] else T["bg_input"])
        self._badge_lbl.configure(bg=self._badge_frame.cget("bg"), fg=T["accent2"])
        self._title_lbl.configure(bg=T["bg_card"], fg=T["fg_primary"])
        self._tag_lbl.configure(bg=T["bg_card"], fg=T["fg_muted"])

        # Scroll area
        self._scroll.configure(fg_color=T["bg_root"],
                                scrollbar_button_color=T["bg_border"],
                                scrollbar_button_hover_color=T["accent"])
        try:
            self._sect_lbl.configure(bg=T["bg_root"], fg=T["fg_muted"])
        except Exception:
            pass

        # Cards
        try:
            for frame, *_ in self._card_frames:
                try:
                    frame._theme_card()
                except Exception:
                    pass
        except Exception:
            pass

        # Status bar
        self._statusbar.configure(bg=T["bg_sidebar"])
        for lbl in [self._sb_left, self._sb_mid, self._sb_right]:
            lbl.configure(bg=T["bg_sidebar"], fg=T["fg_muted"])

        self._draw_strip()

    # ── Lang / Theme toggle ────────────────────────────────────────────────────
    def _toggle_lang(self):
        new = "EN" if self._lang == "TH" else "TH"
        set_lang(new)

    # ── Launch app ─────────────────────────────────────────────────────────────
    def _launch(self, key: str):
        win = self._wins.get(key)
        if win and win.winfo_exists():
            win.lift(); win.focus_force()
            return
        factory = _APP_FACTORIES.get(key)
        if factory:
            win = factory(self)
            self._wins[key] = win


if __name__ == "__main__":
    app = MainLauncher()
    app.mainloop()